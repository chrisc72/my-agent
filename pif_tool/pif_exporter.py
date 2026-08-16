"""PIF Word / PDF 匯出模組"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from database import IngredientDB

CHAPTER_NAMES = {
    1: "產品基本資料",
    2: "完成產品登錄之證明文件",
    3: "全成分名稱及其各別含量",
    4: "產品標籤、仿單、外包裝或容器",
    5: "製造場所符合 GMP 之證明",
    6: "製造方法、流程",
    7: "使用方法、部位、用量、頻率及族群",
    8: "產品使用不良反應資料",
    9: "產品及各別成分之物理及化學特性",
    10: "成分之毒理資料",
    11: "產品安定性試驗報告",
    12: "微生物檢測報告",
    13: "防腐效能試驗報告（PET）",
    14: "功能評估佐證資料",
    15: "與產品接觸之包裝材質資料",
    16: "產品安全資料（安全評估報告）",
}

NUM_ZH = {
    1: "一", 2: "二", 3: "三", 4: "四", 5: "五",
    6: "六", 7: "七", 8: "八", 9: "九", 10: "十",
    11: "十一", 12: "十二", 13: "十三", 14: "十四",
    15: "十五", 16: "十六",
}

# 可上傳附件並在匯出時嵌入內文的章節
ATTACHMENT_CHAPTERS = (2, 4, 5, 6, 11, 12, 13, 14, 15)


def chapter_files(data: dict) -> list[dict]:
    """取出章節附件清單 [{name, path}]，並相容第二章早期的單檔欄位。"""
    files = [f for f in data.get("files", []) if f.get("path")]
    legacy_path = data.get("uploaded_file_path")
    if legacy_path and not any(f["path"] == legacy_path for f in files):
        files.insert(0, {
            "path": legacy_path,
            "name": data.get("uploaded_file_name") or Path(legacy_path).name,
        })
    return files


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_font_kai(run, bold: bool = False):
    run.font.name = "微軟正黑體"
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "微軟正黑體")


def _add_paragraph_kai(doc: Document, text: str, bold: bool = False) -> None:
    """新增段落並套用標楷體。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font_kai(run, bold=bold)


def _add_chapter_heading(doc: Document, num: int, title: str):
    p = doc.add_heading(f"第{NUM_ZH[num]}章　{title}", level=1)
    run = p.runs[0]
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x9B)
    _set_font_kai(run, bold=True)


def _embed_uploaded_file(doc: Document, file_path: str) -> None:
    """將上傳的 PDF（逐頁）或圖片嵌入 Word 文件。路徑不存在時靜默略過。"""
    from io import BytesIO

    p = Path(file_path)
    if not p.exists():
        return

    suffix = p.suffix.lower()
    if suffix in (".jpg", ".jpeg", ".png"):
        doc.add_picture(str(p), width=Inches(6))
    elif suffix == ".pdf":
        import fitz  # pymupdf
        pdf = fitz.open(str(p))
        for page in pdf:
            pix = page.get_pixmap(dpi=150)
            doc.add_picture(BytesIO(pix.tobytes("png")), width=Inches(6))
        pdf.close()


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = ""
        run_k = cells[0].paragraphs[0].add_run(k)
        _set_font_kai(run_k, bold=True)
        _set_cell_bg(cells[0], "E8F0FE")
        cells[1].text = ""
        run_v = cells[1].paragraphs[0].add_run(v or "（待補）")
        _set_font_kai(run_v)
    doc.add_paragraph()


def export_word(pif_id: int, db: IngredientDB, output_path: str):
    doc = Document()

    # 文件標題
    pif_doc = db.get_pif_document(pif_id)
    if not pif_doc:
        raise ValueError(f"找不到 PIF 文件 ID={pif_id}")

    title = doc.add_heading("化粧品產品資訊檔案", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        _set_font_kai(title.runs[0], bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Product Information File\n{pif_doc['product_name']}")
    sub_run.font.size = Pt(13)
    _set_font_kai(sub_run)

    doc.add_page_break()

    ingredients = db.get_pif_ingredients(pif_id)

    for num in range(1, 17):
        _add_chapter_heading(doc, num, CHAPTER_NAMES[num])

        if num == 1:
            ch1 = db.get_chapter_data(pif_id, 1)
            ch1_rows = [("中文品名", ch1.get("product_name", ""))]
            if ch1.get("product_name_en"):
                ch1_rows.append(("英文品名", ch1.get("product_name_en", "")))
            ch1_rows += [
                ("產品種類", ch1.get("product_category", "")),
                ("產品劑型", ch1.get("formulation", "")),
                ("產品用途", ch1.get("product_use", "")),
                ("製造作業場所",
                 f"廠名：{ch1.get('mfg_name','')}\n地址：{ch1.get('mfg_address','')}\n國別：{ch1.get('mfg_country','')}"),
                ("包裝作業場所",
                 f"廠名：{ch1.get('pkg_name','')}\n地址：{ch1.get('pkg_address','')}\n國別：{ch1.get('pkg_country','')}"),
                ("產品製造業者",
                 f"公司負責人：{ch1.get('responsible_person','')}\n"
                 f"產品責任業者：{ch1.get('responsible_party','')}\n"
                 f"地址：{ch1.get('responsible_address','')}\n"
                 f"聯絡電話：{ch1.get('contact','')}"),
            ]
            _add_kv_table(doc, ch1_rows)

        elif num == 3:
            if ingredients:
                headers = ["序號", "INCI 名稱", "CAS No.", "含量 W/W%"]
                table = doc.add_table(rows=1 + len(ingredients), cols=4)
                table.style = "Table Grid"
                for j, h in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = ""
                    run_h = cell.paragraphs[0].add_run(h)
                    run_h.bold = True
                    run_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_font_kai(run_h, bold=True)
                    _set_cell_bg(cell, "1A569B")
                for i, ing in enumerate(ingredients):
                    row = table.rows[i + 1]
                    values = [
                        str(i + 1),
                        ing["inci_name"],
                        ing.get("cas_number") or "",
                        f"{ing.get('percentage'):.5f}" if ing.get("percentage") is not None else "",
                    ]
                    for j, val in enumerate(values):
                        cell = row.cells[j]
                        cell.text = ""
                        _set_font_kai(cell.paragraphs[0].add_run(val))
            else:
                _add_paragraph_kai(doc, "（尚未填入成分清單）")

        elif num == 10:
            ch10 = db.get_chapter_data(pif_id, 10)
            if ch10.get("results"):
                _add_paragraph_kai(doc, f"產品類型：{ch10.get('product_type', '')}　有效日暴露量：{ch10.get('effective_daily_g', '')} g/day")
                headers10 = ["INCI 名稱", "含量%", "評估依據", "NOAEL (mg/kg/day)", "DAp%", "SED (mg/kg/day)", "MoS", "結論"]
                table10 = doc.add_table(rows=1 + len(ch10["results"]), cols=8)
                table10.style = "Table Grid"
                for j, h in enumerate(headers10):
                    cell = table10.rows[0].cells[j]
                    cell.text = ""
                    run_h = cell.paragraphs[0].add_run(h)
                    run_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_font_kai(run_h, bold=True)
                    _set_cell_bg(cell, "1A569B")
                for i, r in enumerate(ch10["results"]):
                    row10 = table10.rows[i + 1]
                    vals10 = [
                        r.get("INCI 名稱", ""), str(r.get("含量%", "")), str(r.get("評估依據", "—")),
                        str(r.get("NOAEL (mg/kg/day)", "—")), str(r.get("DAp%", "")),
                        str(r.get("SED (mg/kg/day)", "—")), str(r.get("MoS", "—")), str(r.get("結論", "")),
                    ]
                    for j, val in enumerate(vals10):
                        cell = row10.cells[j]
                        cell.text = ""
                        _set_font_kai(cell.paragraphs[0].add_run(val))
            else:
                _add_paragraph_kai(doc, "（尚未完成 SED/MoS 計算）")

        elif num == 16:
            _add_paragraph_kai(doc, "安全性評估匯總表", bold=True)
            _add_paragraph_kai(doc, "MoS = PoDsys / SED = (PoD × Bioavailability) / (Exposure × Dermal Absorption)；門檻 MoS ≥ 100")
            ch16_10 = db.get_chapter_data(pif_id, 10)
            if ch16_10.get("results"):
                headers16 = ["成分名稱", "含量w/w%", "Exposure", "皮膚吸收%", "SED",
                             "PoD", "ToE", "Bioavail%", "MoS", "RCR(MoE)", "評估方式"]
                keys16 = ["INCI 名稱", "含量%", "Exposure (mg/kg bw/day)", "DAp%", "SED (mg/kg/day)",
                          "PoD (mg/kg bw/day)", "ToE value", "Bioavailability%", "MoS", "RCR(MoE)", "評估依據"]
                table16 = doc.add_table(rows=1 + len(ch16_10["results"]), cols=len(headers16))
                table16.style = "Table Grid"
                for j, h in enumerate(headers16):
                    cell = table16.rows[0].cells[j]
                    cell.text = ""
                    run_h = cell.paragraphs[0].add_run(h)
                    run_h.font.size = Pt(8)
                    run_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_font_kai(run_h, bold=True)
                    _set_cell_bg(cell, "1A569B")
                for i, r in enumerate(ch16_10["results"]):
                    row16 = table16.rows[i + 1]
                    for j, key in enumerate(keys16):
                        cell = row16.cells[j]
                        cell.text = ""
                        run_v = cell.paragraphs[0].add_run(str(r.get(key, "—")))
                        run_v.font.size = Pt(8)
                        _set_font_kai(run_v)
            else:
                _add_paragraph_kai(doc, "（尚未完成第十章 SED/MoS 計算，暫無匯總表）")
            doc.add_paragraph()
            for line in db.get_chapter_data(pif_id, 16).get("content", "").split("\n"):
                _add_paragraph_kai(doc, line)

        else:
            data = db.get_chapter_data(pif_id, num)
            content = data.get("content", "（待補）")
            for line in content.split("\n"):
                _add_paragraph_kai(doc, line)
            if num in ATTACHMENT_CHAPTERS:
                for f in chapter_files(data):
                    doc.add_paragraph()
                    if Path(f["path"]).exists():
                        _add_paragraph_kai(doc, f"附件：{f['name']}", bold=True)
                        _embed_uploaded_file(doc, f["path"])
                    else:
                        _add_paragraph_kai(doc, f"附件：{f['name']}（檔案遺失，請重新上傳）", bold=True)

        doc.add_paragraph()

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def export_pdf(docx_path: str, pdf_path: str) -> str:
    """將 Word 檔轉換為 PDF。優先用 docx2pdf（需要 Word），備用 LibreOffice。"""
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return pdf_path
    except ImportError:
        pass

    import subprocess
    import shutil
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        out_dir = str(Path(pdf_path).parent)
        subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                        "--outdir", out_dir, docx_path], check=True)
        return pdf_path

    raise RuntimeError(
        "無法匯出 PDF：請安裝 docx2pdf（pip install docx2pdf）或 LibreOffice。"
    )


def export_inci_word(pif_id: int, db: IngredientDB) -> tuple[bytes, str]:
    """匯出精簡版全成分 Word，包含 INCI 成分表與成分標示文字。回傳 (docx_bytes, filename)。"""
    from io import BytesIO

    pif_doc = db.get_pif_document(pif_id)
    product_name = pif_doc["product_name"] if pif_doc else f"PIF_{pif_id}"
    filename = f"{product_name} 全成分.docx"

    doc = Document()

    title_para = doc.add_heading(f"{product_name} 全成分", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ingredients = db.get_pif_ingredients(pif_id)

    _add_paragraph_kai(doc, "INCI 成分表（由高至低）", bold=True)
    if ingredients:
        headers = ["序號", "INCI 名稱", "CAS No.", "含量 W/W%"]
        table = doc.add_table(rows=1 + len(ingredients), cols=4)
        table.style = "Table Grid"
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            run_h = cell.paragraphs[0].add_run(h)
            run_h.bold = True
            run_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_font_kai(run_h, bold=True)
            _set_cell_bg(cell, "1A569B")
        for i, ing in enumerate(ingredients):
            row = table.rows[i + 1]
            values = [
                str(i + 1),
                ing["inci_name"],
                ing.get("cas_number") or "",
                f"{ing['percentage']:.5f}" if ing.get("percentage") is not None else "",
            ]
            for j, val in enumerate(values):
                cell = row.cells[j]
                cell.text = ""
                _set_font_kai(cell.paragraphs[0].add_run(val))
    else:
        _add_paragraph_kai(doc, "（尚未填入成分清單）")

    doc.add_paragraph()
    _add_paragraph_kai(doc, "成分標示文字（由高至低）", bold=True)
    inci_text = ", ".join(ing["inci_name"] for ing in ingredients) if ingredients else ""
    _add_paragraph_kai(doc, inci_text)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue(), filename


def export_ingredient_breakdown_excel(
    formula_rows: list[dict], db, product_name: str = ""
) -> tuple[bytes, str]:
    """逐原料拆解 INCI 對照表，格式參照衛福部申報用 Excel 範本。"""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from io import BytesIO

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "原料編號對照表"

    center = Alignment(horizontal="center", vertical="center")

    if product_name:
        ws.append([f"{product_name} 原料編號對照表"])
        for cell in ws[ws.max_row]:
            cell.alignment = center
        ws.append([])

    header = ["原料編號", "成分名稱 (INCI)", "CAS No."]
    ws.append(header)
    header_row = ws.max_row
    for col in range(1, 4):
        cell = ws.cell(row=header_row, column=col)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1A569B")
        cell.alignment = center

    seen_codes: set[str] = set()
    sorted_rows = sorted(
        formula_rows,
        key=lambda r: str(r.get("原料編號", "")).strip().upper(),
    )
    for row in sorted_rows:
        code = str(row.get("原料編號", "")).strip()
        if not code or code in seen_codes:
            continue
        seen_codes.add(code)
        material = db.get_raw_material_by_code(code)
        if material is None:
            ws.append([code, "(原料資料庫無此編號)", ""])
            for cell in ws[ws.max_row]:
                cell.alignment = center
            continue
        components = db.get_components_by_material_id(material["id"])
        if not components:
            ws.append([code, material.get("product_name", ""), ""])
            for cell in ws[ws.max_row]:
                cell.alignment = center
            continue
        first = True
        for comp in components:
            ws.append([
                code if first else "",
                comp.get("inci_name", ""),
                comp.get("cas_number", "") or "",
            ])
            for cell in ws[ws.max_row]:
                cell.alignment = center
            first = False

    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 42
    ws.column_dimensions["C"].width = 16

    buf = BytesIO()
    wb.save(buf)
    filename = f"{product_name} 原料編號對照表.xlsx" if product_name else "原料編號對照表.xlsx"
    return buf.getvalue(), filename


def _build_coa_sds_index_pdf(
    product_name: str,
    sorted_mids: list,
    materials: dict,
    inci_by_material: dict,
    page_map: dict,
) -> bytes:
    """用 python-docx 建立 COA/SDS/IFRA 索引頁，轉為 PDF bytes。"""
    import tempfile
    import os
    from io import BytesIO

    doc = Document()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_p.add_run(product_name)
    _set_font_kai(t_run, bold=True)
    t_run.font.size = Pt(14)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub_p.add_run("原料安全資料（COA / SDS / IFRA）附錄索引")
    _set_font_kai(s_run, bold=True)
    s_run.font.size = Pt(11)

    doc.add_paragraph()

    headers = ["序號", "原料編號", "原料名稱", "所含 INCI 成分", "文件類型", "頁碼"]
    col_widths = [
        Inches(0.35), Inches(0.75), Inches(1.3),
        Inches(2.8), Inches(0.75), Inches(0.55),
    ]

    table = doc.add_table(rows=1 + len(sorted_mids), cols=6)
    table.style = "Table Grid"

    for j, (h, w) in enumerate(zip(headers, col_widths)):
        cell = table.rows[0].cells[j]
        cell.width = w
        cell.text = ""
        rh = cell.paragraphs[0].add_run(h)
        rh.bold = True
        rh.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_font_kai(rh, bold=True)
        rh.font.size = Pt(9)
        _set_cell_bg(cell, "1A569B")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for seq, mid in enumerate(sorted_mids, 1):
        minfo = materials[mid]
        row = table.rows[seq]
        incis = inci_by_material.get(mid, [])
        inci_text = "\n".join(incis) if incis else "(無 INCI 對應)"

        file_types = sorted(set(f["file_type"] for f in minfo["files"]))
        has_pdf_data = mid in page_map

        if has_pdf_data:
            sp, ep = page_map[mid]
            page_str = str(sp) if sp == ep else f"{sp}–{ep}"
            seq_str = str(seq)
        else:
            page_str = "—"
            seq_str = f"⚠ {seq}"
            file_types = []

        values = [
            seq_str,
            minfo["ingredient_code"],
            minfo["material_name"],
            inci_text,
            ", ".join(file_types) if file_types else "無原始文件",
            page_str,
        ]

        for j, (val, w) in enumerate(zip(values, col_widths)):
            cell = row.cells[j]
            cell.width = w
            cell.text = ""
            rv = cell.paragraphs[0].add_run(val)
            _set_font_kai(rv)
            rv.font.size = Pt(8)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "index.docx")
        pdf_path = os.path.join(tmpdir, "index.pdf")
        doc.save(docx_path)
        export_pdf(docx_path, pdf_path)
        with open(pdf_path, "rb") as f:
            return f.read()


def export_coa_sds_appendix(pif_id: int, db: IngredientDB) -> tuple[bytes, str]:
    """生成原料 COA/SDS/IFRA 合併附錄 PDF。
    第一頁為索引表（INCI→原料→頁碼），後續頁為實際 COA/SDS/IFRA 內容。
    回傳 (pdf_bytes, filename)。
    """
    import fitz
    import tempfile
    import os
    from io import BytesIO

    pif_doc = db.get_pif_document(pif_id)
    product_name = pif_doc["product_name"] if pif_doc else f"PIF_{pif_id}"

    inci_rows = db.get_pif_material_inci_map(pif_id)
    file_rows = db.get_pif_material_coa_sds(pif_id)

    # INCI → material grouping
    inci_by_material: dict[int, list[str]] = {}
    for row in inci_rows:
        inci_by_material.setdefault(row["material_id"], []).append(row["inci_name"])

    # material_id → {ingredient_code, material_name, files}
    materials: dict[int, dict] = {}
    for row in file_rows:
        mid = row["material_id"]
        if mid not in materials:
            materials[mid] = {
                "ingredient_code": row["ingredient_code"] or "",
                "material_name": row["material_name"] or "",
                "files": [],
            }
        if row["file_id"] is not None and row["file_data"] is not None:
            materials[mid]["files"].append({
                "filename": row["filename"] or "",
                "file_type": row["file_type"] or "",
                "file_data": row["file_data"],
            })

    sorted_mids = sorted(materials.keys(), key=lambda m: materials[m]["ingredient_code"])

    # Pass 1: merge COA/SDS/IFRA PDFs and track page numbers
    body_doc = fitz.open()
    page_map: dict[int, tuple[int, int]] = {}  # material_id → (start, end) 1-indexed after index page

    for mid in sorted_mids:
        minfo = materials[mid]
        start = body_doc.page_count + 2  # page 1 = index
        inserted = False
        for f in minfo["files"]:
            fname = f["filename"].lower()
            try:
                if fname.endswith(".pdf"):
                    src = fitz.open(stream=f["file_data"], filetype="pdf")
                    body_doc.insert_pdf(src)
                    src.close()
                    inserted = True
                elif fname.endswith(".docx"):
                    with tempfile.TemporaryDirectory() as tmpdir:
                        dp = os.path.join(tmpdir, "src.docx")
                        pp = os.path.join(tmpdir, "src.pdf")
                        with open(dp, "wb") as fp:
                            fp.write(f["file_data"])
                        export_pdf(dp, pp)
                        src = fitz.open(pp)
                        body_doc.insert_pdf(src)
                        src.close()
                    inserted = True
            except Exception:
                pass
        end = body_doc.page_count + 1
        if inserted:
            page_map[mid] = (start, end)

    # Build index page PDF
    index_bytes = _build_coa_sds_index_pdf(
        product_name, sorted_mids, materials, inci_by_material, page_map
    )

    # Assemble: index first, then body
    result_doc = fitz.open()
    idx_src = fitz.open(stream=index_bytes, filetype="pdf")
    result_doc.insert_pdf(idx_src)
    idx_src.close()
    result_doc.insert_pdf(body_doc)
    body_doc.close()

    # Stamp page numbers (ASCII, no custom font needed)
    total = result_doc.page_count
    for i, page in enumerate(result_doc):
        label = f"- {i + 1} / {total} -"
        cx = page.rect.width / 2
        page.insert_text(
            fitz.Point(cx - 20, page.rect.height - 12),
            label, fontsize=8, color=(0.45, 0.45, 0.45),
        )

    buf = BytesIO()
    result_doc.save(buf)
    result_doc.close()

    return buf.getvalue(), f"{product_name}_COA_SDS_IFRA附錄.pdf"
