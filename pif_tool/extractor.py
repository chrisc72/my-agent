"""從 PDF / Word / Excel 提取文字內容"""
import io
import fitz  # PyMuPDF
from docx import Document
import openpyxl

PDF_MAX_CHARS = 15_000
_TOX_KEYWORDS = [
    "conclusion", "noael", "loael", "mos", "margin of safety",
    "safe", "concentration", "exposure", "toxicolog",
    "dermal", "systemic", "sensitiz", "acceptable",
]


def extract_text_from_file(uploaded_file) -> str:
    filename = uploaded_file.name.lower()
    content = uploaded_file.read()

    if filename.endswith(".pdf"):
        return _from_pdf(content)
    elif filename.endswith(".docx"):
        return _from_docx(content)
    elif filename.endswith((".xlsx", ".xls")):
        return _from_excel(content)
    elif filename.endswith(".md"):
        return content.decode("utf-8", errors="replace")
    return ""


def _from_pdf(content: bytes) -> str:
    doc = fitz.open(stream=content, filetype="pdf")
    all_pages = [page.get_text() for page in doc]
    full_text = "\n".join(all_pages)

    if len(full_text) <= PDF_MAX_CHARS:
        return full_text

    # 優先取含毒理關鍵詞的頁面（結論、NOAEL 等在後段）
    key_pages = [p for p in all_pages if any(kw in p.lower() for kw in _TOX_KEYWORDS)]
    if key_pages:
        candidate = "\n---\n".join(key_pages)
        if len(candidate) <= PDF_MAX_CHARS:
            return candidate
        # 結論在後段：前 1/4 背景 + 後 3/4（含 Conclusion 章節）
        front = PDF_MAX_CHARS // 4
        rear = PDF_MAX_CHARS - front
        return candidate[:front] + "\n...[中段省略]...\n" + candidate[-rear:]

    # 前 1/3 + 後 2/3（SCCS 結論通常在文件後段）
    front = PDF_MAX_CHARS // 3
    rear = PDF_MAX_CHARS - front
    return full_text[:front] + "\n...[中段省略]...\n" + full_text[-rear:]


def _from_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _from_excel(content: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== 工作表: {sheet_name} ===")
        for row in ws.iter_rows():
            cells = []
            for cell in row:
                val = cell.value
                if val is not None and isinstance(val, (int, float)) and "%" in (cell.number_format or ""):
                    val = val * 100
                cells.append(str(val) if val is not None else "")
            row_text = "\t".join(cells)
            if row_text.strip().replace("\t", ""):
                parts.append(row_text)
    return "\n".join(parts)
