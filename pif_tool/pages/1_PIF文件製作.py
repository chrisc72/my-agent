"""PIF 文件製作頁面 — 16 章節填寫、SED/MoS 計算、Word/PDF 匯出"""
import re
import shutil
import sys
import time
from pathlib import Path

import anthropic as _anthropic
import pandas as pd
import streamlit as st

sys.path.insert(0, str(Path(__file__).parent.parent))
from database import IngredientDB
from pif_exporter import ATTACHMENT_CHAPTERS, chapter_files
from tw_regulations import lookup_tw_regulation
from tw_food_ingredients import lookup_food_ingredient


def _claude_create(**kwargs):
    """呼叫 Claude API，遇到 529 Overloaded 時指數退避重試最多 5 次。"""
    client = _anthropic.Anthropic()
    waits = [5, 15, 30, 60, 120]
    for attempt in range(5):
        try:
            return client.messages.create(**kwargs)
        except _anthropic.APIStatusError as e:
            if e.status_code == 529 and attempt < 4:
                time.sleep(waits[attempt])
                continue
            raise


# ─── 常數 ─────────────────────────────────────────────────────────────────────

CHAPTER_NAMES = {
    1: "產品基本資料",
    2: "完成產品登錄之證明文件",
    3: "全成分名稱及其各別含量",
    4: "產品標籤、仿單、外包裝或容器",
    5: "製造場所符合 GMP 之證明",
    6: "製造方法、流程",
    7: "使用方法、部位、用量、頻率及族群",
    8: "產品使用不良反應資料",
    9: "產品及各別成分之物理及化學特性",
    10: "成分之毒理資料",
    11: "產品安定性試驗報告",
    12: "微生物檢測報告",
    13: "防腐效能試驗報告（PET）",
    14: "功能評估佐證資料",
    15: "與產品接觸之包裝材質資料",
    16: "產品安全資料（安全評估報告）",
}

CHAPTER_HINTS = {
    1: "包含產品中英文名稱、類別（一般/特定用途化粧品）、劑型、主要功能宣稱、製造廠名稱地址、公司統編等。",
    2: "填入衛福部食藥署化粧品產品登錄系統的確認文件編號與登錄日期，並可附上截圖或確認信。",
    3: "（本章由上方表格管理）",
    4: "核對外包裝/容器/標籤/仿單是否齊全，並確認符合第7條標示規定（品名、廠商地址、全成分、內容量、保存期限等）。可於下方上傳標籤/包裝圖檔，匯出時會併入本章節。",
    5: "提供 GMP 證書號碼與有效期限，或填寫自我聲明書。尚未取得證書者可使用業者自我聲明。可於下方上傳證書 PDF 或圖片，匯出時會併入本章節。",
    6: "描述製造流程概述，如原料投入順序、溫度控制、混合時間、充填條件等。建議附上批次生產標準書（SOP）。可於下方上傳流程圖或 SOP，匯出時會併入本章節。",
    7: "說明使用部位、建議用量、使用頻率、駐留時間（留置型/沖洗型）、適用族群（成人/嬰幼兒）及不適用族群與警語。",
    8: "建立不良反應管理機制，上市前可填「尚無不良反應紀錄」。嚴重不良反應需於 15 日內向食藥署通報。",
    9: "提供成品的外觀、顏色、氣味、pH、黏度、比重等規格，以及主要活性成分的物化特性（分子量、溶解度等）。",
    10: "（本章由下方 SED/MoS 計算器自動產生）",
    11: "提供加速試驗（40°C/75%RH，6個月）與長期試驗（25°C/60%RH，12個月）的穩定性結果。如免除需由安全評估人員說明理由。可於下方上傳試驗報告，匯出時會併入本章節。",
    12: "提供總生菌數、大腸桿菌、金黃色葡萄球菌、綠膿桿菌、黴菌及酵母菌的檢驗報告。眼部/嬰幼兒產品需更嚴格規格。可於下方上傳檢驗報告，匯出時會併入本章節。",
    13: "防腐效能試驗（ISO 11930:2019 或 USP <51>）報告。微生物低風險產品符合特定條件者可豁免。可於下方上傳試驗報告，匯出時會併入本章節。",
    14: "功能宣稱（如保濕、美白）需有佐證：文獻、體外評估報告、消費者實測報告或成分含量檢驗報告。可於下方上傳佐證資料，匯出時會併入本章節。",
    15: "列出所有與產品接觸的包裝材質（瓶身、瓶蓋、泵頭等），說明材質、規格、供應商及相容性評估結果。可於下方上傳材質證明，匯出時會併入本章節。",
    16: "由符合資格的安全資料簽署人員（醫/藥/毒理/化粧品相關科系）完成最終安全評估並簽署。",
}

UPLOAD_ROOT = Path(__file__).parent.parent / "data" / "pif_uploads"

# content_json 中存放檔案路徑的鍵
_PATH_KEYS = ("uploaded_file_path", "formula_file_path")


def _safe_dir_component(s: str) -> str:
    """去掉 Windows 檔名非法字元，保留中文與空格，結尾不留點或空白。"""
    s = re.sub(r'[\\/:*?"<>|]', "_", s or "").strip().rstrip(".")
    return s or "未命名"


def _safe_file_name(s: str) -> str:
    return re.sub(r"[^\w.\-]", "_", s)


def _pif_dir_name(pif_doc: dict) -> str:
    """{貨品編號}_{品名}，無編號時只用品名。"""
    name = _safe_dir_component(pif_doc.get("product_name", ""))
    doc_num = _safe_dir_component(pif_doc["document_number"]) if pif_doc.get("document_number") else ""
    return f"{doc_num}_{name}" if doc_num else name


def pif_upload_dir(db: IngredientDB, pif_doc: dict) -> Path:
    """該 PIF 的附件目錄。不建立目錄。"""
    dirname = _pif_dir_name(pif_doc)
    # 兩份 PIF 算出同名目錄時，讓 id 較大的那份加後綴，結果對每份 PIF 都是穩定的
    clashes = [
        d["id"] for d in db.get_pif_documents()
        if d["id"] != pif_doc["id"] and _pif_dir_name(d) == dirname
    ]
    if clashes and min(clashes) < pif_doc["id"]:
        dirname = f"{dirname}_pif{pif_doc['id']}"
    return UPLOAD_ROOT / dirname


def ensure_pif_upload_dir(db: IngredientDB, pif_doc: dict) -> Path:
    d = pif_upload_dir(db, pif_doc)
    d.mkdir(parents=True, exist_ok=True)
    return d


def _rewrite_chapter_paths(db: IngredientDB, pif_id: int, old_dir: Path, new_dir: Path) -> int:
    """把該 PIF 章節資料裡指向 old_dir 的檔案路徑改指 new_dir。回傳更新的章節數。"""
    changed = 0
    for ch in range(1, 17):
        data = db.get_chapter_data(pif_id, ch)
        if not data:
            continue
        touched = False

        def _remap(p: str) -> str:
            nonlocal touched
            try:
                rel = Path(p).relative_to(old_dir)
            except ValueError:
                return p
            touched = True
            return str(new_dir / rel)

        for key in _PATH_KEYS:
            if data.get(key):
                data[key] = _remap(data[key])
        for f in data.get("files", []):
            if f.get("path"):
                f["path"] = _remap(f["path"])

        if touched:
            db.save_chapter_data(pif_id, ch, data)
            changed += 1
    return changed


def _delete_pif_with_files(db: IngredientDB, pif_doc: dict):
    """刪除 PIF 及其附件目錄，避免留下孤兒檔。"""
    shutil.rmtree(pif_upload_dir(db, pif_doc), ignore_errors=True)
    # 保險：清掉可能殘留在根目錄的舊平鋪檔
    if UPLOAD_ROOT.exists():
        for p in UPLOAD_ROOT.glob(f"pif{pif_doc['id']}_*"):
            p.unlink(missing_ok=True)
    db.delete_pif_document(pif_doc["id"])


def _save_chapter_files(db: IngredientDB, pif_id: int, chapter_num: int, files: list[dict]):
    """寫回附件清單。重新讀取避免蓋掉其他欄位，並清除第二章的舊單檔欄位。"""
    data = db.get_chapter_data(pif_id, chapter_num)
    if files:
        data["files"] = files
    else:
        data.pop("files", None)  # 留下空清單會讓章節完成度誤判為已完成
    data.pop("uploaded_file_path", None)
    data.pop("uploaded_file_name", None)
    db.save_chapter_data(pif_id, chapter_num, data)


def render_chapter_attachments(
    db: IngredientDB, pif_doc: dict, chapter_num: int,
    label: str = "上傳附件（PDF 或圖片，可多選）",
) -> list[dict]:
    """多檔上傳 + 已上傳清單 + 逐檔刪除。回傳目前的附件清單。"""
    pif_id = pif_doc["id"]
    files = chapter_files(db.get_chapter_data(pif_id, chapter_num))

    ver_key = f"ch{chapter_num}_upl_ver_{pif_id}"
    ver = st.session_state.get(ver_key, 0)
    uploaded = st.file_uploader(
        label,
        type=["pdf", "jpg", "jpeg", "png"],
        accept_multiple_files=True,
        key=f"ch{chapter_num}_uploader_{pif_id}_{ver}",
    )

    if uploaded:
        existing = {f["name"] for f in files}
        added = False
        for uf in uploaded:
            if uf.name in existing:
                continue
            upload_dir = ensure_pif_upload_dir(db, pif_doc)
            save_path = upload_dir / f"ch{chapter_num:02d}_{_safe_file_name(uf.name)}"
            save_path.write_bytes(uf.getvalue())
            files.append({"path": str(save_path), "name": uf.name})
            existing.add(uf.name)
            added = True
        if added:
            _save_chapter_files(db, pif_id, chapter_num, files)

    if files:
        st.caption(f"已上傳 {len(files)} 個附件，匯出時會依序嵌入本章節")
        for i, f in enumerate(files):
            col_name, col_del = st.columns([6, 1])
            if Path(f["path"]).exists():
                col_name.write(f"📎 {f['name']}")
            else:
                col_name.write(f"⚠️ {f['name']}（檔案遺失，請重新上傳）")
            if col_del.button("🗑️", key=f"del_ch{chapter_num}_{pif_id}_{i}", help="刪除此附件"):
                Path(f["path"]).unlink(missing_ok=True)
                _save_chapter_files(db, pif_id, chapter_num, [x for j, x in enumerate(files) if j != i])
                st.session_state[ver_key] = ver + 1
                st.rerun()
    else:
        st.info("尚未上傳附件")

    return files

# SCCS 標準日暴露量表（來源：SCCS Notes of Guidance 12th revision, 2023）
SCCS_EXPOSURE: dict[str, dict] = {
    "身體乳 (Body lotion)": {"daily_g": 7.82, "retention": 1.0},
    "面霜 (Face cream)": {"daily_g": 1.54, "retention": 1.0},
    "潤手霜 (Hand cream)": {"daily_g": 2.16, "retention": 1.0},
    "洗浴凝膠 (Shower gel)": {"daily_g": 18.67, "retention": 0.01},
    "洗手皂 (Hand wash soap)": {"daily_g": 20.00, "retention": 0.01},
    "洗髮精 (Shampoo)": {"daily_g": 10.46, "retention": 0.01},
    "頭髮調理劑 (Hair conditioner)": {"daily_g": 3.92, "retention": 0.01},
    "頭髮造型產品 (Hair styling)": {"daily_g": 4.00, "retention": 0.10},
    "液態底粧 (Liquid foundation)": {"daily_g": 0.51, "retention": 1.0},
    "卸粧產品 (Makeup remover)": {"daily_g": 5.00, "retention": 0.10},
    "眼影 (Eye shadow)": {"daily_g": 0.02, "retention": 1.0},
    "睫毛膏 (Mascara)": {"daily_g": 0.025, "retention": 1.0},
    "眼線膏 (Eye liner)": {"daily_g": 0.005, "retention": 1.0},
    "唇膏 / 護唇膏 (Lip products)": {"daily_g": 0.057, "retention": 1.0},
    "除臭劑（非噴霧）": {"daily_g": 1.50, "retention": 1.0},
    "除臭劑（噴霧，乙醇）": {"daily_g": 1.43, "retention": 1.0},
    "除臭劑（噴霧，非乙醇）": {"daily_g": 0.69, "retention": 1.0},
    "牙膏（成人）": {"daily_g": 2.75, "retention": 0.05},
    "漱口水": {"daily_g": 21.62, "retention": 0.10},
    "自行輸入暴露量": {"daily_g": 0.0, "retention": 1.0},
}

# TTC（Threshold of Toxicological Concern）備援值，mg/kg bw/day
# 來源：KB「SCCS 全身暴露劑量與安全邊際計算方法」；僅在無 NOAEL 時作輔助，不得作為唯一結論依據
TTC_VALUES: dict[str, float] = {
    "Class I": 0.0491,
    "Class II": 0.0029,    # 資料庫支持不足，比照 Class III
    "Class III": 0.0029,
    "基因毒性警訊": 0.0000025,
}

# ─── 工具函式 ──────────────────────────────────────────────────────────────────

def _parse_num(raw: str) -> float | None:
    try:
        return float(raw.replace(",", ""))
    except (ValueError, AttributeError):
        return None


def extract_noael(text: str) -> float | None:
    """從毒理摘要文字中提取 NOAEL（mg/kg bw/day）。

    優先取 AI 標註、用於 MoS 的「關鍵 NOAEL」，其次取帶 bw/day 的全身性 NOAEL，
    最後才退回寬鬆比對，避免盲抓文件中第一個出現的數值。
    """
    if not text:
        return None

    # 1) AI 明確標註、用於 MoS 的關鍵 NOAEL
    m = re.search(r"關鍵\s*NOAEL[^0-9無\n]*?([\d,.]+)\s*mg\s*/?\s*kg", text, re.IGNORECASE)
    if m:
        v = _parse_num(m.group(1))
        if v is not None:
            return v
    # 若 AI 標註「關鍵 NOAEL … 無」，代表文獻無可用 NOAEL，不再往下臆測
    if re.search(r"關鍵\s*NOAEL[^0-9\n]{0,20}無", text):
        return None

    # 2) 帶 bw/day 的全身性 NOAEL（SCCS/CIR 計算 MoS 的口徑）
    for pat in [
        r"NOAEL[^\n]{0,80}?([\d,.]+)\s*mg\s*/?\s*kg\s*(?:bw|body\s*weight)?\s*/?\s*(?:bw)?\s*(?:day|d)\b",
        r"([\d,.]+)\s*mg\s*/?\s*kg\s*bw\s*/?\s*day[^\n]{0,40}?NOAEL",
    ]:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            v = _parse_num(m.group(1))
            if v is not None:
                return v

    # 3) 寬鬆備援
    for pat in [
        r"NOAEL\s*[：:=]\s*([\d,.]+)\s*mg\s*/?\s*kg",
        r"NOAEL\s+(?:of\s+)?([\d,.]+)\s*mg\s*/?\s*kg",
        r"no.{0,20}observed.{0,20}adverse.{0,20}effect.{0,20}level\s*[：:=]?\s*([\d,.]+)\s*mg\s*/?\s*kg",
    ]:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            v = _parse_num(m.group(1))
            if v is not None:
                return v
    return None


def extract_safe_conc(text: str) -> float | None:
    """從 SCCS/CIR 文獻文字中提取建議安全使用濃度（%）。"""
    if not text:
        return None
    patterns = [
        # SCCS/CIR 標準用語：「maximum (use) concentration of X%」（可含 "of up to"）
        r"maximum\s+(?:use\s+)?concentration\s+of\s+(?:up\s+to\s+)?([\d.]+)\s*%",
        # 「safe ... at/up to a maximum of X%」
        r"safe\b[^.%]{0,80}?(?:up\s+to|at)\s+(?:a\s+maximum\s+of\s+)?([\d.]+)\s*%",
        # 「concentrations (of) up to X%」
        r"concentrations?\s+(?:of\s+)?up\s+to\s+([\d.]+)\s*%",
        # 「concentrations of X% ... safe/acceptable」
        r"concentrations?\s+of\s+([\d.]+)\s*%[^.]{0,40}(?:safe|acceptable|no\s+(?:safety\s+)?(?:concern|risk))",
        r"safe\s+(?:for\s+use\s+)?(?:in\s+cosmetics?\s+)?(?:at|when\s+used\s+at|up\s+to)\s+(?:concentrations?\s+(?:of\s+|up\s+to\s+))?([\d.]+)\s*%",
        r"concentrations?\s+(?:of\s+|up\s+to\s+)?([\d.]+)\s*%\s+(?:is\s+|are\s+)?safe",
        r"(?:not\s+exceed|no\s+more\s+than)\s+(?:a\s+concentration\s+of\s+)?([\d.]+)\s*%",
        # 中文（AI 毒理摘要輸出）
        r"(?:建議安全使用量|安全使用量|安全上限|安全使用濃度)\s*[：:]\s*([\d.]+)\s*%",
        r"(?:最大|最高)(?:使用)?濃度\s*(?:為|：|:|\s)?\s*([\d.]+)\s*%",
        r"at\s+([\d.]+)\s*%[^,\.]{0,30}(?:safe|acceptable|no\s+risk)",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            try:
                val = float(m.group(1))
                if 0 < val <= 100:
                    return val
            except ValueError:
                pass
    return None


def parse_tw_limit_pct(limit_str: str) -> float | None:
    """從台灣法規限量字串中提取最保守（最低）的百分比數值。"""
    if not limit_str:
        return None
    matches = re.findall(r"([\d.]+)%", limit_str)
    if not matches:
        return None
    return min(float(m) for m in matches)


def calc_exposure(effective_daily_g: float, concentration_pct: float) -> float:
    """該成分每日外用暴露量（尚未乘皮膚吸收率）mg/kg bw/day。"""
    return effective_daily_g * (concentration_pct / 100) * 1000 / 60


def calc_sed(effective_daily_g: float, concentration_pct: float, dap_pct: float = 100.0) -> float:
    """計算 SED（全身暴露量）。
    SED (mg/kg bw/day) = 有效日用量(g) × 成分濃度/100 × 皮膚吸收率/100 × 1000 ÷ 60
    """
    return calc_exposure(effective_daily_g, concentration_pct) * (dap_pct / 100)


def calc_mos(pod: float | None, sed: float, bioavailability_pct: float = 100.0) -> float | None:
    if pod is None or sed <= 0:
        return None
    return pod * (bioavailability_pct / 100) / sed


def chapter_emoji(has_data: bool) -> str:
    return "✅" if has_data else "⚠️"


@st.cache_resource
def get_db() -> IngredientDB:
    db_path = str(Path(__file__).parent.parent / "data" / "ingredients.db")
    return IngredientDB(db_path)


# ─── AI 輔助工具 ─────────────────────────────────────────────────────────────

def extract_ch1_from_doc(file_bytes: bytes, filename: str) -> dict:
    """從產品登錄文件（PDF/圖片）提取第一章欄位，透過 Claude 視覺模型解析。"""
    import json as _json
    import base64
    import anthropic as _anthropic

    ext = Path(filename).suffix.lower().lstrip(".")
    if ext == "pdf":
        media_type, block_type = "application/pdf", "document"
    elif ext in ("jpg", "jpeg"):
        media_type, block_type = "image/jpeg", "image"
    else:
        media_type, block_type = "image/png", "image"

    b64 = base64.standard_b64encode(file_bytes).decode()
    prompt = (
        "你是台灣化粧品法規專家。請從這份化粧品產品登錄確認文件中提取以下資訊，"
        "直接回傳 JSON（不要有其他說明）：\n\n"
        '{"product_name":"","product_name_en":"","product_category":"",'
        '"formulation":"","product_use":"",'
        '"mfg_name":"","mfg_address":"","mfg_country":"台灣",'
        '"pkg_name":"","pkg_address":"","pkg_country":"台灣",'
        '"responsible_person":"","responsible_party":"","responsible_address":"",'
        '"contact":"",'
        '"registration_number":"","registration_date":""}\n\n'
        "提取規則：\n"
        "- product_name：中文品名\n"
        "- product_name_en：英文品名，找不到填空字串\n"
        "- product_category：直接複製文件原文（如「洗髮用化粧品類」），不要解讀或分類\n"
        "- product_use：產品用途（如「清潔頭髮及頭皮」）\n"
        "- mfg_name/mfg_address/mfg_country：製造作業場所的廠名、地址、國別\n"
        "- pkg_name/pkg_address/pkg_country：包裝作業場所的廠名、地址、國別（如與製造相同則填相同值）\n"
        "- responsible_party：從「案件資訊」區塊下方表格的「登錄業者名稱」欄位取得\n"
        "- contact：從「案件資訊」區塊下方表格的「電話」欄位取得\n"
        "- responsible_address：若文件有「地址」欄則填入，否則填空字串\n"
        "- responsible_person：文件中通常沒有此資訊，填空字串（由使用者手動補填）\n"
        "- registration_number / registration_date：登錄確認文件編號與日期\n"
        "找不到的欄位填空字串。"
    )
    response = _claude_create(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        messages=[{
            "role": "user",
            "content": [
                {"type": block_type, "source": {"type": "base64", "media_type": media_type, "data": b64}},
                {"type": "text", "text": prompt},
            ],
        }],
    )
    text = response.content[0].text.strip()
    try:
        return _json.loads(text)
    except Exception:
        pass
    m = re.search(r"```(?:json)?\s*(\{[\s\S]*?\})\s*```", text)
    if m:
        try:
            return _json.loads(m.group(1))
        except Exception:
            pass
    m = re.search(r"\{[\s\S]*\}", text)
    if m:
        try:
            return _json.loads(m.group(0))
        except Exception:
            pass
    return {}


# ─── 第一章：產品基本資料 ──────────────────────────────────────────────────────

def render_chapter_1(db: IngredientDB, pif_id: int):
    st.subheader("一、產品基本資料")
    data = db.get_chapter_data(pif_id, 1)

    # AI 自動填寫提示與按鈕
    ch2 = db.get_chapter_data(pif_id, 2)
    ch2_file = next((f for f in chapter_files(ch2) if Path(f["path"]).exists()), None)
    if data.get("ai_filled"):
        st.success("✅ 本章欄位已由 AI 依第二章登錄文件自動填寫，請確認內容後點選「儲存」。")
    elif ch2_file:
        st.info("💡 第二章已有上傳文件，可點下方按鈕讓 AI 自動提取並填入欄位。")
    else:
        st.info("💡 可先到第二章上傳產品登錄確認文件，AI 將自動提取資訊填寫本章。")

    if ch2_file:
        if st.button("🤖 AI 自動填寫（讀取第二章文件）", key="ch1_ai_fill"):
            with st.spinner("AI 正在讀取文件，請稍候..."):
                try:
                    extracted = extract_ch1_from_doc(
                        Path(ch2_file["path"]).read_bytes(),
                        ch2_file["name"],
                    )
                    if extracted:
                        for k, v in extracted.items():
                            if k not in ("registration_number", "registration_date") and v:
                                data[k] = v
                        data["ai_filled"] = True
                        db.save_chapter_data(pif_id, 1, data)
                        st.rerun()
                    else:
                        st.warning("AI 未能提取到資訊，請手動填寫。")
                except Exception as e:
                    st.error(f"AI 分析失敗：{e}")

    # 「同製造作業場所」快速複製按鈕（在 form 外）
    if st.button("📋 包裝作業場所同製造作業場所", key="ch1_copy_pkg"):
        saved = db.get_chapter_data(pif_id, 1)
        saved["pkg_name"] = saved.get("mfg_name", "")
        saved["pkg_address"] = saved.get("mfg_address", "")
        saved["pkg_country"] = saved.get("mfg_country", "台灣")
        db.save_chapter_data(pif_id, 1, saved)
        st.rerun()

    with st.form("ch1_form"):
        st.markdown("#### 基本資料")
        col1, col2 = st.columns(2)
        with col1:
            product_name = st.text_input("中文品名 *", value=data.get("product_name", ""))
            product_category = st.text_input("產品種類", value=data.get("product_category", ""),
                                             placeholder="例：洗髮用化粧品類、護膚用化粧品類")
        with col2:
            product_name_en = st.text_input("英文品名（選填，匯出時顯示於中文品名下方）",
                                            value=data.get("product_name_en", ""))
            formulation = st.text_input("產品劑型", value=data.get("formulation", ""),
                                        placeholder="例：液劑、乳霜、凝膠")
        product_use = st.text_input("產品用途", value=data.get("product_use", ""),
                                    placeholder="例：清潔頭髮及頭皮")

        st.markdown("#### 製造作業場所")
        col3, col4, col5 = st.columns([3, 4, 2])
        with col3:
            mfg_name = st.text_input("廠名", value=data.get("mfg_name", ""), key="mfg_name")
        with col4:
            mfg_address = st.text_input("地址", value=data.get("mfg_address", ""), key="mfg_address")
        with col5:
            mfg_country = st.text_input("國別", value=data.get("mfg_country", "台灣"), key="mfg_country")

        st.markdown("#### 包裝作業場所")
        col6, col7, col8 = st.columns([3, 4, 2])
        with col6:
            pkg_name = st.text_input("廠名", value=data.get("pkg_name", ""), key="pkg_name")
        with col7:
            pkg_address = st.text_input("地址", value=data.get("pkg_address", ""), key="pkg_address")
        with col8:
            pkg_country = st.text_input("國別", value=data.get("pkg_country", "台灣"), key="pkg_country")

        st.markdown("#### 產品製造業者")
        col9, col10 = st.columns(2)
        with col9:
            responsible_person = st.text_input("公司負責人", value=data.get("responsible_person", ""))
            responsible_party = st.text_input("產品責任業者", value=data.get("responsible_party", ""))
        with col10:
            responsible_address = st.text_area("地址", value=data.get("responsible_address", ""), height=90)
        contact = st.text_input("聯絡電話", value=data.get("contact", ""))

        if st.form_submit_button("💾 儲存第一章", type="primary"):
            db.save_chapter_data(pif_id, 1, {
                "product_name": product_name,
                "product_name_en": product_name_en,
                "product_category": product_category,
                "formulation": formulation,
                "product_use": product_use,
                "mfg_name": mfg_name,
                "mfg_address": mfg_address,
                "mfg_country": mfg_country,
                "pkg_name": pkg_name,
                "pkg_address": pkg_address,
                "pkg_country": pkg_country,
                "responsible_person": responsible_person,
                "responsible_party": responsible_party,
                "responsible_address": responsible_address,
                "contact": contact,
                "ai_filled": data.get("ai_filled", False),
            })
            db.update_pif_document(pif_id, {
                "product_type": product_category,
                "formulation_type": formulation,
                "claims": product_use,
                "manufacturer": f"{mfg_name}\n{mfg_address}",
                "responsible_party": responsible_party,
            })
            st.success("第一章已儲存！")
            st.rerun()


# ─── 第二章：產品登錄之證明文件 ─────────────────────────────────────────────────

def render_chapter_2(db: IngredientDB, pif_doc: dict):
    pif_id = pif_doc["id"]
    st.subheader("二、完成產品登錄之證明文件")
    st.caption("上傳衛福部食藥署化粧品產品登錄確認文件，AI 可自動讀取並填寫第一章基本資料。")

    files = render_chapter_attachments(
        db, pif_doc, 2, label="上傳登錄確認文件（PDF 或圖片，可多選）"
    )
    data = db.get_chapter_data(pif_id, 2)

    primary = next((f for f in files if Path(f["path"]).exists()), None)
    if primary:
        if st.button("🤖 AI 分析文件並自動填寫第一章", type="primary", key=f"ch2_ai_fill_{pif_id}"):
            with st.spinner(f"AI 正在讀取「{primary['name']}」，請稍候..."):
                try:
                    extracted = extract_ch1_from_doc(Path(primary["path"]).read_bytes(), primary["name"])
                    if extracted:
                        ch1 = db.get_chapter_data(pif_id, 1)
                        for k, v in extracted.items():
                            if k not in ("registration_number", "registration_date") and v:
                                ch1[k] = v
                        ch1["ai_filled"] = True
                        db.save_chapter_data(pif_id, 1, ch1)
                        if extracted.get("registration_number"):
                            data["registration_number"] = extracted["registration_number"]
                        if extracted.get("registration_date"):
                            data["registration_date"] = extracted["registration_date"]
                        db.save_chapter_data(pif_id, 2, data)
                        st.success("✅ 第一章已自動填寫！請切換到第一章確認後儲存。")
                        st.json({k: v for k, v in extracted.items() if v})
                    else:
                        st.warning("AI 未能從文件中提取到資訊，請手動填寫第一章。")
                except Exception as e:
                    st.error(f"AI 分析失敗：{e}")
        if len(files) > 1:
            st.caption(f"AI 只會讀取第一個附件「{primary['name']}」。")

    st.divider()
    st.markdown("#### 登錄資訊")
    col1, col2 = st.columns(2)
    with col1:
        reg_num = st.text_input("登錄確認文件編號", value=data.get("registration_number", ""), key=f"ch2_regnum_{pif_id}")
    with col2:
        reg_date = st.text_input("登錄日期", value=data.get("registration_date", ""), placeholder="例：2024-01-15", key=f"ch2_regdate_{pif_id}")
    notes = st.text_area("其他備註", value=data.get("content", ""), height=120, key=f"ch2_notes_{pif_id}")

    if st.button("💾 儲存第二章", type="primary", key=f"save_ch2_{pif_id}"):
        data.update({"registration_number": reg_num, "registration_date": reg_date, "content": notes})
        db.save_chapter_data(pif_id, 2, data)
        st.success("第二章已儲存！")
        st.rerun()


# ─── AI 提取第三章成分 ───────────────────────────────────────────────────────────

def extract_ch3_from_doc(file_bytes: bytes, filename: str) -> list[dict]:
    """從配方文件（PDF/Word/Excel）提取全成分表，透過 Claude 解析。"""
    import json as _json
    import base64
    import io
    import anthropic as _anthropic

    ext = Path(filename).suffix.lower().lstrip(".")

    prompt = (
        "你是台灣化粧品法規專家。請從這份文件中提取全成分表，"
        "直接回傳 JSON 陣列（不要有其他說明）：\n\n"
        '[{"inci_name":"成分名稱","cas_number":"CAS號或空字串","percentage":含量數字或null}]\n\n'
        "提取規則：\n"
        "- inci_name：INCI 名稱或成分名稱（必填）\n"
        "- cas_number：CAS 號碼，找不到填空字串\n"
        "- percentage：含量（W/W%），純數字，找不到填 null\n"
        "只包含成分行，忽略標題行和空白行。"
    )

    if ext == "pdf":
        b64 = base64.standard_b64encode(file_bytes).decode()
        messages = [{
            "role": "user",
            "content": [
                {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": b64}},
                {"type": "text", "text": prompt},
            ],
        }]
    elif ext == "docx":
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        text = "\n".join(parts)
        messages = [{"role": "user", "content": f"配方文件內容：\n\n{text}\n\n{prompt}"}]
    elif ext in ("xlsx", "xls"):
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== 工作表: {sheet_name} ===")
            for row in ws.iter_rows():
                cells = []
                for cell in row:
                    val = cell.value
                    if val is not None and isinstance(val, (int, float)) and "%" in (cell.number_format or ""):
                        val = val * 100
                    cells.append(str(val) if val is not None else "")
                row_text = "\t".join(cells)
                if row_text.strip().replace("\t", ""):
                    parts.append(row_text)
        text = "\n".join(parts)
        messages = [{"role": "user", "content": f"配方文件內容：\n\n{text}\n\n{prompt}"}]
    else:
        return []

    response = _claude_create(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        messages=messages,
    )
    raw = response.content[0].text.strip()

    try:
        return _json.loads(raw)
    except Exception:
        pass
    m = re.search(r"```(?:json)?\s*(\[[\s\S]*?\])\s*```", raw)
    if m:
        try:
            return _json.loads(m.group(1))
        except Exception:
            pass
    m = re.search(r"\[[\s\S]*\]", raw)
    if m:
        try:
            return _json.loads(m.group(0))
        except Exception:
            pass
    return []


# ─── 第三章：複方展開 helpers ────────────────────────────────────────────────

def _resolve_comp_percentage(comp: dict) -> float | None:
    """取成分百分比：有精確值優先，否則嘗試從 percentage_range 取中間值。"""
    if comp.get("percentage") is not None:
        try:
            return float(comp["percentage"])
        except (TypeError, ValueError):
            pass
    pr = comp.get("percentage_range") or ""
    if pr:
        parts = re.findall(r"[\d.]+", str(pr))
        if len(parts) >= 2:
            return (float(parts[0]) + float(parts[-1])) / 2
        if len(parts) == 1:
            return float(parts[0])
    return None


def expand_formula_to_inci(formula_rows: list[dict], db) -> tuple[list[dict], list[str]]:
    """
    formula_rows: [{"原料編號": str, "含量 W/W%": float}, ...]
    Returns: (inci_list, warnings)
      inci_list: [{"inci_name", "cas_number", "percentage"}, ...] 按含量由高至低排序
    """
    inci_totals: dict[str, dict] = {}
    warnings: list[str] = []
    for row in formula_rows:
        code = str(row.get("原料編號", "")).strip()
        usage_pct = row.get("含量 W/W%") or 0
        if not code or not usage_pct:
            continue
        material = db.get_raw_material_by_code(code)
        if material is None:
            warnings.append(f"找不到原料編號「{code}」，請確認原料清單資料庫已有此編號。")
            continue
        components = db.get_components_by_material_id(material["id"])
        resolved = [(c, _resolve_comp_percentage(c)) for c in components]
        valid = [(c, p) for c, p in resolved if p is not None]
        estimated = [c for c, p in resolved if p is not None and c.get("percentage") is None]
        if not valid:
            warnings.append(f"原料「{code} {material['product_name']}」在資料庫中無成分組成資料（缺少含量數值）。")
            continue
        if estimated:
            names = "、".join(c["inci_name"] for c in estimated)
            warnings.append(f"注意：原料「{code}」中的 {names} 使用範圍值估算（取中間值），計算結果僅供參考。")
        for comp, pct in valid:
            inci_in_product = (pct / 100) * usage_pct
            key = comp["inci_name"]
            if key in inci_totals:
                inci_totals[key]["percentage"] += inci_in_product
            else:
                inci_totals[key] = {
                    "inci_name": comp["inci_name"],
                    "cas_number": comp.get("cas_number") or "",
                    "percentage": inci_in_product,
                }
    result = sorted(inci_totals.values(), key=lambda x: x["percentage"], reverse=True)
    for item in result:
        item["percentage"] = round(item["percentage"], 4)
    return result, warnings


def extract_formula_from_doc(file_bytes: bytes, filename: str) -> list[dict]:
    """從配方文件（Excel/Word/PDF）提取原料編號＋使用含量，透過 Claude 解析。"""
    import json as _json
    import base64
    import io
    import anthropic as _anthropic

    ext = Path(filename).suffix.lower().lstrip(".")
    prompt = (
        "你是化粧品配方管理專家。請從這份文件中提取配方的原料清單，"
        "直接回傳 JSON 陣列（不要有其他說明）：\n\n"
        '[{"ingredient_code":"原料編號","usage_pct":含量數字}]\n\n'
        "提取規則：\n"
        "- ingredient_code：原料編號或代號（必填），通常是英數字組合如 RM-001、A001 等\n"
        "- usage_pct：此原料在成品中的使用含量（W/W%），純數字\n"
        "只包含原料行，忽略標題行、合計行和空白行。"
    )

    if ext == "pdf":
        b64 = base64.standard_b64encode(file_bytes).decode()
        messages = [{"role": "user", "content": [
            {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": b64}},
            {"type": "text", "text": prompt},
        ]}]
    elif ext == "docx":
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(c.text.strip() for c in row.cells))
        messages = [{"role": "user", "content": f"配方文件：\n\n{chr(10).join(parts)}\n\n{prompt}"}]
    elif ext in ("xlsx", "xls"):
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                cells = []
                for cell in row:
                    val = cell.value
                    if val is not None and isinstance(val, (int, float)) and "%" in (cell.number_format or ""):
                        val = val * 100
                    cells.append(str(val) if val is not None else "")
                row_text = "\t".join(cells)
                if row_text.strip().replace("\t", ""):
                    parts.append(row_text)
        messages = [{"role": "user", "content": f"配方文件：\n\n{chr(10).join(parts)}\n\n{prompt}"}]
    else:
        return []

    response = _claude_create(
        model="claude-sonnet-4-6", max_tokens=2000, messages=messages
    )
    raw = response.content[0].text.strip()
    for pattern in [r"```(?:json)?\s*(\[[\s\S]*?\])\s*```", r"\[[\s\S]*\]"]:
        m = re.search(pattern, raw)
        if m:
            try:
                group = m.group(1) if m.lastindex and m.lastindex >= 1 else m.group(0)
                return _json.loads(group)
            except Exception:
                pass
    try:
        return _json.loads(raw)
    except Exception:
        return []


# ─── 第三章：全成分清單（核心） ────────────────────────────────────────────────

def render_chapter_3(db: IngredientDB, pif_doc: dict):
    pif_id = pif_doc["id"]
    st.subheader("三、全成分名稱及其各別含量")

    # ── 兩種輸入方式：上傳 INCI 成分表 vs 從原料編號展開 ──
    data_ch3 = db.get_chapter_data(pif_id, 3)
    tab_formula, tab_inci = st.tabs(["🔢 用原料編號建立（系統展開 INCI）", "📄 已有 INCI 清單，直接匯入"])

    # ── Tab A：上傳含 INCI 的配方文件 ──────────────────────────────
    with tab_inci:
        st.caption("適用於文件裡「已經列出 INCI 成分名稱」的情況（Word / Excel / PDF），AI 直接讀取 INCI 名稱、CAS 號與含量並填入下方成分表，不需查原料庫、不需展開。")

        saved_path = data_ch3.get("uploaded_file_path", "")
        saved_name = data_ch3.get("uploaded_file_name", "")

        col_up, col_st = st.columns([2, 1])
        with col_up:
            uploaded = st.file_uploader(
                "上傳配方文件（Word .docx、Excel .xlsx、PDF .pdf）",
                type=["docx", "xlsx", "xls", "pdf"],
                key=f"ch3_uploader_{pif_id}",
            )
        with col_st:
            if saved_name:
                st.success(f"已上傳：{saved_name}")
            else:
                st.info("尚未上傳文件")

        if uploaded is not None:
            upload_dir = ensure_pif_upload_dir(db, pif_doc)
            save_path = upload_dir / f"ch03_{_safe_file_name(uploaded.name)}"
            save_path.write_bytes(uploaded.getvalue())
            data_ch3["uploaded_file_path"] = str(save_path)
            data_ch3["uploaded_file_name"] = uploaded.name
            db.save_chapter_data(pif_id, 3, data_ch3)
            saved_path = str(save_path)
            saved_name = uploaded.name

        if saved_path and Path(saved_path).exists():
            _existing_count = len(db.get_pif_ingredients(pif_id))
            if _existing_count > 0:
                st.warning(f"⚠️ 目前已有 {_existing_count} 筆成分資料，重新提取將完全覆蓋。")
            if st.button("🤖 AI 自動提取成分資料", type="primary", key=f"ch3_ai_fill_{pif_id}"):
                with st.spinner("AI 正在讀取文件，請稍候..."):
                    try:
                        extracted = extract_ch3_from_doc(Path(saved_path).read_bytes(), saved_name)
                        if extracted:
                            ingredients = []
                            for item in extracted:
                                inci = str(item.get("inci_name", "")).strip()
                                if not inci:
                                    continue
                                cas = str(item.get("cas_number", "")).strip()
                                pct = item.get("percentage")
                                comp = db.lookup_component_by_inci_cas(inci, cas)
                                ingredients.append({
                                    "inci_name": inci,
                                    "cas_number": cas,
                                    "percentage": float(pct) if pct is not None else None,
                                    "function": "",
                                    "component_id": comp["id"] if comp else None,
                                })
                            db.save_pif_ingredients(pif_id, ingredients)
                            matched = sum(1 for i in ingredients if i["component_id"])
                            st.success(f"✅ AI 已提取 {len(ingredients)} 個成分（✅ 有毒理資料：{matched} 個），請確認後視需要修改再儲存。")
                            st.rerun()
                        else:
                            st.warning("AI 未能從文件中提取到成分資料，請手動填寫。")
                    except Exception as e:
                        st.error(f"AI 分析失敗：{e}")

    # ── Tab B：從原料編號展開 INCI ──────────────────────────────────
    with tab_formula:
        st.caption("提供每支原料的「原料編號」與「在成品中的使用含量（%）」，系統從原料資料庫查出各原料的 INCI 組成，計算各成分在成品中的實際比例，相同 INCI 自動加總，最後由高至低排序。")
        st.caption("👇 以下兩種只是「輸入原料編號」的方式，擇一使用即可：上傳檔案自動辨識，或手動逐筆輸入。")

        sub_upload, sub_manual = st.tabs(["📄 上傳配方文件（AI 解析原料編號）", "✏️ 手動輸入原料編號"])

        formula_key = f"ch3_formula_{pif_id}"
        if formula_key not in st.session_state:
            _saved_rows = data_ch3.get("formula_rows", [])
            if _saved_rows:
                st.session_state[formula_key] = pd.DataFrame(_saved_rows)
            else:
                st.session_state[formula_key] = pd.DataFrame(columns=["原料編號", "含量 W/W%"])

        with sub_upload:
            st.caption("上傳含有「原料編號」與「使用含量」欄位的配方文件，AI 自動辨識並填入下方原料表格。")

            _formula_saved_name = data_ch3.get("formula_file_name", "")
            _formula_saved_rows = data_ch3.get("formula_rows", [])
            col_fup, col_fst = st.columns([2, 1])
            with col_fst:
                if _formula_saved_name and _formula_saved_rows:
                    st.success(f"已上傳：{_formula_saved_name}（{len(_formula_saved_rows)} 筆）")
                else:
                    st.info("尚未上傳配方文件")
            with col_fup:
                formula_file = st.file_uploader(
                    "上傳配方文件",
                    type=["xlsx", "xls", "docx", "pdf"],
                    key=f"ch3_formula_file_{pif_id}",
                )
            if formula_file and st.button("🤖 AI 解析原料編號", key=f"ch3_parse_formula_{pif_id}"):
                with st.spinner("AI 正在讀取配方文件..."):
                    try:
                        parsed = extract_formula_from_doc(formula_file.getvalue(), formula_file.name)
                        if parsed:
                            _parsed_rows = [
                                {"原料編號": r.get("ingredient_code", ""), "含量 W/W%": r.get("usage_pct", 0)}
                                for r in parsed
                            ]
                            st.session_state[formula_key] = pd.DataFrame(_parsed_rows)

                            upload_dir = ensure_pif_upload_dir(db, pif_doc)
                            formula_save_path = upload_dir / f"ch03_formula_{_safe_file_name(formula_file.name)}"
                            formula_save_path.write_bytes(formula_file.getvalue())

                            data_ch3["formula_file_path"] = str(formula_save_path)
                            data_ch3["formula_file_name"] = formula_file.name
                            data_ch3["formula_rows"] = _parsed_rows
                            db.save_chapter_data(pif_id, 3, data_ch3)

                            st.success(f"AI 解析完成，共 {len(parsed)} 筆原料，請確認後點「計算 INCI 成分表」。")
                            st.rerun()
                        else:
                            st.warning("AI 未能從文件中解析到原料編號，請改用手動輸入。")
                    except _anthropic.APIStatusError as e:
                        if e.status_code == 529:
                            st.error("Claude API 目前流量過高（服務暫時過載），請稍等 1-2 分鐘後再按一次「AI 解析原料編號」。")
                        else:
                            st.error(f"AI 解析失敗：{e}")
                    except Exception as e:
                        st.error(f"AI 解析失敗：{e}")

        with sub_manual:
            st.caption("直接在下方表格中輸入原料編號與使用含量。")

        formula_df = st.data_editor(
            st.session_state[formula_key],
            num_rows="dynamic",
            use_container_width=True,
            column_config={
                "原料編號": st.column_config.TextColumn("原料編號 *", required=True),
                "含量 W/W%": st.column_config.NumberColumn("含量 W/W%", min_value=0.0, max_value=100.0, format="%.4f"),
            },
            key=f"ch3_formula_editor_{pif_id}",
        )
        st.session_state[formula_key] = formula_df

    # 此產品配方實際用到的原料編號，用來把成分比對限縮在這些原料內，
    # 避免同 INCI 名稱（如多支香精都叫 Fragrance）跨原料配錯。
    formula_codes: list[str] = []
    if "原料編號" in formula_df.columns:
        formula_codes = [
            str(c).strip()
            for c in formula_df["原料編號"].dropna()
            if str(c).strip()
        ]

    with tab_formula:
        from pif_exporter import export_ingredient_breakdown_excel
        _pif_doc = db.get_pif_document(pif_id)
        _pname = _pif_doc["product_name"] if _pif_doc else ""
        _rows_for_dl = formula_df.dropna(subset=["原料編號"]).to_dict("records")

        _col_calc, _col_dl = st.columns([1, 1])
        with _col_calc:
            if st.button("⚗️ 計算 INCI 成分表", type="primary", key=f"ch3_calc_{pif_id}"):
                rows = formula_df.dropna(subset=["原料編號"]).to_dict("records")
                with st.spinner("正在計算 INCI 成分表..."):
                    inci_result, warns = expand_formula_to_inci(rows, db)
                st.session_state[f"ch3_inci_result_{pif_id}"] = inci_result
                for w in warns:
                    st.warning(w)
                if not inci_result:
                    st.warning("未產生任何 INCI 資料，請確認原料編號與含量是否正確填寫。")
        with _col_dl:
            if _rows_for_dl:
                _xl_bytes, _xl_fname = export_ingredient_breakdown_excel(_rows_for_dl, db, _pname)
                st.download_button(
                    "📊 下載原料編號對照表 (Excel for 衛福部)",
                    data=_xl_bytes,
                    file_name=_xl_fname,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key=f"ch3_dl_breakdown_{pif_id}",
                )

        result_key = f"ch3_inci_result_{pif_id}"
        if result_key in st.session_state and st.session_state[result_key]:
            inci_result = st.session_state[result_key]
            _c_title, _c_recalc, _c_clear = st.columns([4, 1, 1])
            with _c_title:
                st.markdown("**INCI 成分表（由高至低）**")
            with _c_recalc:
                if st.button("🔄 重新計算", key=f"ch3_recalc_{pif_id}"):
                    rows = formula_df.dropna(subset=["原料編號"]).to_dict("records")
                    with st.spinner("計算中..."):
                        new_result, warns = expand_formula_to_inci(rows, db)
                    st.session_state[result_key] = new_result
                    st.session_state.pop(f"ch3_inci_text_{pif_id}", None)
                    for w in warns:
                        st.warning(w)
                    st.rerun()
            with _c_clear:
                if st.button("🗑️ 清除", key=f"ch3_clear_inci_{pif_id}"):
                    del st.session_state[result_key]
                    st.session_state.pop(f"ch3_inci_text_{pif_id}", None)
                    st.rerun()
            st.dataframe(
                pd.DataFrame(inci_result).rename(columns={
                    "inci_name": "INCI 名稱",
                    "cas_number": "CAS No.",
                    "percentage": "含量 W/W%",
                }),
                use_container_width=True,
                hide_index=True,
            )
            inci_text = ", ".join(item["inci_name"] for item in inci_result)
            st.markdown("**成分標示文字（由高至低，複製貼上用）**")
            st.text_area(
                "",
                value=inci_text,
                height=100,
                key=f"ch3_inci_text_{pif_id}",
                label_visibility="collapsed",
            )
            if st.button("📋 套用至下方成分表", key=f"ch3_apply_{pif_id}"):
                ingredients = []
                for item in inci_result:
                    comp = db.lookup_component_by_inci_cas(item["inci_name"], item["cas_number"], formula_codes)
                    ingredients.append({
                        "inci_name": item["inci_name"],
                        "cas_number": item["cas_number"],
                        "percentage": item["percentage"],
                        "function": "",
                        "component_id": comp["id"] if comp else None,
                    })
                db.save_pif_ingredients(pif_id, ingredients)
                matched = sum(1 for i in ingredients if i["component_id"])
                st.success(f"已套用 {len(ingredients)} 個 INCI 至成分表。✅ 有毒理資料：{matched} 個")
                st.rerun()


    st.divider()

    # ── 成分編輯表格 ──
    existing = db.get_pif_ingredients(pif_id)
    if existing:
        df_init = pd.DataFrame([{
            "INCI 名稱": r["inci_name"],
            "CAS No.": r["cas_number"] or "",
            "含量 W/W%": r["percentage"] if r["percentage"] is not None else 0.0,
        } for r in existing])
    else:
        df_init = pd.DataFrame(columns=["INCI 名稱", "CAS No.", "含量 W/W%"])

    edited_df = st.data_editor(
        df_init,
        num_rows="dynamic",
        use_container_width=True,
        column_config={
            "INCI 名稱": st.column_config.TextColumn("INCI 名稱 *", required=True),
            "CAS No.": st.column_config.TextColumn("CAS No."),
            "含量 W/W%": st.column_config.NumberColumn("含量 W/W%", min_value=0.0, max_value=100.0, format="%.5f"),
        },
        key=f"ch3_editor_{pif_id}",
    )

    total = edited_df["含量 W/W%"].sum() if "含量 W/W%" in edited_df.columns else 0
    if abs(total - 100.0) < 0.01:
        st.success(f"總計：{total:.5f}% ✅")
    elif total > 0:
        st.warning(f"總計：{total:.5f}%（應為 100%）")

    col_save, col_reset, col_dl, col_info = st.columns([1, 1, 1, 1])
    with col_save:
        save_clicked = st.button("💾 儲存第三章並比對資料庫", type="primary")
    with col_reset:
        if existing and st.button("🗑️ 重置成分表", type="secondary", key=f"ch3_reset_{pif_id}"):
            st.session_state[f"ch3_confirm_reset_{pif_id}"] = True
    with col_dl:
        if existing:
            from pif_exporter import export_inci_word
            docx_bytes, dl_filename = export_inci_word(pif_id, db)
            st.download_button(
                "📥 匯出全成分 Word",
                data=docx_bytes,
                file_name=dl_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"ch3_dl_word_{pif_id}",
            )
    if st.session_state.get(f"ch3_confirm_reset_{pif_id}", False):
        st.warning("⚠️ 確定要清空所有成分資料嗎？此操作無法復原。")
        col_yes, col_no, _ = st.columns([1, 1, 4])
        with col_yes:
            if st.button("✅ 確認清空", key=f"ch3_reset_yes_{pif_id}"):
                db.clear_pif_chapter3(pif_id)
                st.session_state.pop(f"ch3_confirm_reset_{pif_id}", None)
                st.success("成分表已清空。")
                st.rerun()
        with col_no:
            if st.button("❌ 取消", key=f"ch3_reset_no_{pif_id}"):
                st.session_state.pop(f"ch3_confirm_reset_{pif_id}", None)
                st.rerun()

    if save_clicked:
        rows = edited_df.dropna(subset=["INCI 名稱"]).to_dict("records")
        ingredients = []
        matched, unmatched = 0, 0
        for row in rows:
            inci = str(row.get("INCI 名稱", "")).strip()
            cas = str(row.get("CAS No.", "")).strip()
            pct = row.get("含量 W/W%")
            if not inci:
                continue
            comp = db.lookup_component_by_inci_cas(inci, cas, formula_codes)
            comp_id = comp["id"] if comp else None
            if comp_id:
                matched += 1
            else:
                unmatched += 1
            ingredients.append({
                "inci_name": inci,
                "cas_number": cas,
                "percentage": float(pct) if pct is not None else None,
                "function": "",
                "component_id": comp_id,
            })
        db.save_pif_ingredients(pif_id, ingredients)
        st.success(f"已儲存 {len(ingredients)} 個成分。✅ 有毒理資料：{matched} 個，⚠️ 需補充：{unmatched} 個")
        st.rerun()

    # 顯示比對狀態
    if existing:
        st.divider()
        st.caption("目前比對狀態：")
        status_data = [{
            "成分": r["inci_name"],
            "含量%": r["percentage"],
            "毒理資料": "✅ 有" if r.get("component_id") else "⚠️ 需補充",
        } for r in existing]
        st.dataframe(pd.DataFrame(status_data), use_container_width=True, hide_index=True)
        if any(r.get("component_id") is None for r in existing):
            st.info("⚠️ 部分成分尚未建立毒理資料。請至「原料分析」頁面為這些成分新增原料資料後，再回來重新儲存第三章。")


# ─── 第十章：SED / MoS 計算器（核心） ─────────────────────────────────────────

def render_chapter_10(db: IngredientDB, pif_id: int):
    st.subheader("十、成分之毒理資料")
    st.caption("依 SCCS Notes of Guidance（12th revision, 2023）計算各成分的 SED 與 MoS。MoS ≥ 100 為可接受。")

    ingredients = db.get_pif_ingredients(pif_id)
    if not ingredients:
        st.warning("⚠️ 請先填寫第三章（全成分清單）並儲存，才能進行毒理評估。")
        return

    st.markdown("#### 產品類型設定")
    col1, col2 = st.columns(2)
    with col1:
        product_type_sel = st.selectbox(
            "選擇產品類型（帶入 SCCS 標準暴露量）",
            list(SCCS_EXPOSURE.keys()),
            key="ch10_product_type",
        )
    exp_data = SCCS_EXPOSURE[product_type_sel]
    with col2:
        if product_type_sel == "自行輸入暴露量":
            custom_daily = st.number_input("每日用量 g/day", min_value=0.0, value=1.0, step=0.1, key="ch10_custom_daily")
            custom_retention = st.number_input("駐留因子（0.01=沖洗型，1.0=留置型）", min_value=0.0, max_value=1.0, value=1.0, step=0.01, key="ch10_custom_ret")
            effective_daily_g = custom_daily * custom_retention
        else:
            effective_daily_g = exp_data["daily_g"] * exp_data["retention"]
            st.metric("有效日暴露量（已乘駐留因子）", f"{effective_daily_g:.4f} g/day",
                      help=f"原始日用量 {exp_data['daily_g']} g × 駐留因子 {exp_data['retention']}")

    st.markdown("#### 各成分 NOAEL 與皮膚吸收率輸入")
    st.caption("系統已嘗試從資料庫提取 NOAEL，請確認數值或手動填入。DAp 預設 100%（保守估算）。")

    tox_rows = []
    for ing in ingredients:
        comp = None
        if ing.get("component_id"):
            comp = db.get_component_by_id(ing["component_id"])

        noael_from_db = None
        if comp:
            noael_from_db = extract_noael(comp.get("toxicology_summary", "") or "")
            if not noael_from_db:
                noael_from_db = extract_noael(comp.get("sccs_data", "") or "")

        tox_rows.append({
            "INCI 名稱": ing["inci_name"],
            "含量%": ing["percentage"] or 0.0,
            "NOAEL (mg/kg/day)": ing.get("noael_manual") or noael_from_db or None,
            "DAp%": ing.get("dap_override") or 100.0,
            "Bioavailability%": ing.get("bioavailability_override") or 50.0,
            "Cramer Class": ing.get("cramer_class") or "",
            "毒理資料": "✅ 資料庫" if comp else "⚠️ 需補充",
        })

    tox_df = pd.DataFrame(tox_rows)
    edited_tox = st.data_editor(
        tox_df,
        use_container_width=True,
        column_config={
            "INCI 名稱": st.column_config.TextColumn("INCI 名稱", disabled=True),
            "含量%": st.column_config.NumberColumn("含量%", disabled=True, format="%.5f"),
            "NOAEL (mg/kg/day)": st.column_config.NumberColumn("NOAEL (mg/kg/day)", min_value=0.0, format="%.2f",
                                                                  help="從文獻或資料庫取得的 NOAEL 值，可手動修改"),
            "DAp%": st.column_config.NumberColumn("DAp%（皮膚吸收率）", min_value=0.0, max_value=100.0, format="%.0f",
                                                   help="預設 100%（保守估算）"),
            "Bioavailability%": st.column_config.NumberColumn("Bioavailability%（口服吸收率）", min_value=0.0, max_value=100.0, format="%.0f",
                                                               help="無實測時 SCCS 預設 50%"),
            "Cramer Class": st.column_config.SelectboxColumn("Cramer 分級（無 NOAEL 時採 TTC）",
                                                             options=["", "Class I", "Class II", "Class III", "基因毒性警訊", "不需評估"],
                                                             help="無 NOAEL 時依結構判定 TTC；惰性成分（如水）選「不需評估」"),
            "毒理資料": st.column_config.TextColumn("資料來源", disabled=True),
        },
        key=f"ch10_tox_editor_{pif_id}",
    )

    if st.button("🔬 計算 SED 與 MoS", type="primary", key="ch10_calc"):
        results = []
        for _, row in edited_tox.iterrows():
            inci = row["INCI 名稱"]
            conc = row["含量%"]
            noael = row["NOAEL (mg/kg/day)"]
            dap = row["DAp%"]
            bioavail = row["Bioavailability%"]
            cramer = row["Cramer Class"]

            # 找到對應的成分記錄（用於 fallback 資料查詢）
            ing_rec = next((i for i in ingredients if i["inci_name"] == inci), None)
            cas = ing_rec.get("cas_number", "") if ing_rec else ""
            comp_rec = db.get_component_by_id(ing_rec["component_id"]) if ing_rec and ing_rec.get("component_id") else None

            # 先計算 Exposure / SED（純物理量，與 NOAEL 無關），再依 PoD 來源算 MoS
            noael_val = float(noael) if pd.notna(noael) and noael else None
            exposure = calc_exposure(effective_daily_g, conc) if conc else 0.0
            sed = exposure * (dap / 100)

            if cramer == "不需評估":
                pod_used, pod_source, mos = None, "不需評估", None
            elif noael_val:
                pod_used, pod_source = noael_val, "NOAEL"
                mos = calc_mos(noael_val, sed, bioavail)
            elif cramer in TTC_VALUES:
                pod_used, pod_source = TTC_VALUES[cramer], f"TTC（Cramer {cramer}）"
                mos = calc_mos(pod_used, sed, 100.0)   # TTC 已是全身可接受劑量，不再乘生體可用率
            else:
                pod_used, pod_source, mos = None, None, None
            rcr = (1 / mos) if mos else None

            # 候選一：台灣法規上限（精確比對，避免子字串誤判）
            tw_entries_all = lookup_tw_regulation(inci_name=inci, cas_number=cas)
            inci_lower = inci.lower().strip()
            cas_clean = cas.strip()
            tw_entries = [
                e for e in tw_entries_all
                if any(n.lower().strip() == inci_lower for n in e["inci_names"])
                or (cas_clean and any(c.strip() == cas_clean for c in e.get("cas_numbers", [])))
            ]
            safe_conc_tw = None
            tw_limit_display = ""
            for entry in tw_entries:
                pct = parse_tw_limit_pct(entry["limit"])
                if pct is not None:
                    if safe_conc_tw is None or pct < safe_conc_tw:
                        safe_conc_tw = pct
                        tw_limit_display = entry["limit"]

            # 候選一之二：食用安全史（WoE）。部位不符者不採計，靜默落回 NOAEL/MoS
            food_res = lookup_food_ingredient(inci)

            # 候選二：文獻建議安全使用濃度（SCCS / CIR / 毒理摘要）
            safe_conc_lit = None
            lit_source = ""
            if comp_rec:
                for field, label in [
                    ("toxicology_summary", "毒理摘要"),
                    ("sccs_data", "SCCS"),
                    ("cir_data", "CIR"),
                ]:
                    val = comp_rec.get(field, "") or ""
                    c = extract_safe_conc(val)
                    if c is not None:
                        safe_conc_lit = c
                        lit_source = label
                        break

            # 判定優先序：不需評估 → 台灣法規上限 → 文獻安全濃度 → 食用安全史 WoE → NOAEL/MoS → TTC 輔助
            if cramer == "不需評估":
                basis = "不需評估"
                conclusion = "惰性成分，不需個別 MoS 評估（As discussion）"

            elif safe_conc_tw is not None:
                basis = f"台灣法規上限（{tw_limit_display}）"
                if conc <= safe_conc_tw:
                    conclusion = f"✅ 安全（添加量 {conc:.4f}% ≤ 台灣法規上限 {safe_conc_tw:.4f}%）"
                else:
                    conclusion = f"🔴 添加量 {conc:.4f}% 超出台灣法規上限 {safe_conc_tw:.4f}%"

            elif safe_conc_lit is not None:
                basis = f"文獻安全濃度（{lit_source}）"
                if conc <= safe_conc_lit:
                    conclusion = f"✅ 安全（添加量 {conc:.4f}% ≤ 文獻安全濃度 {safe_conc_lit:.4f}%）"
                else:
                    conclusion = f"🔴 添加量 {conc:.4f}% 超出文獻建議 {safe_conc_lit:.4f}%，需評估"

            elif food_res["status"] == "food_approved":
                if food_res["part_unspecified"]:
                    parts = "、".join(food_res["allowed_parts"])
                    basis = f"食品原料 WoE（表列部位：{parts}，請確認原料實際使用部位）"
                else:
                    basis = "食品原料 WoE"
                if mos is not None and mos < 100:
                    conclusion = (
                        f"⚠️ 列於台灣「可供食品使用原料一覽表」，但 MoS = {mos:.1f} < 100，需人工複核"
                    )
                else:
                    conclusion = (
                        "✅ 安全（列於台灣「可供食品使用原料一覽表」，具長期歷史食用安全性；"
                        "外用之皮膚吸收低於經口暴露，安全性可接受）"
                    )

            elif pod_source == "NOAEL" and mos is not None:
                basis = "NOAEL / MoS"
                if mos >= 100:
                    conclusion = f"✅ 安全（MoS = {mos:.1f} ≥ 100）"
                else:
                    conclusion = f"🔴 MoS = {mos:.1f} < 100，需進一步評估"

            elif pod_source and pod_source.startswith("TTC") and mos is not None:
                basis = f"TTC（{cramer}，輔助）"
                conclusion = f"⚠️ 無 NOAEL，採 TTC 估算 MoS = {mos:.1f}（僅供輔助，須人工確認）"

            else:
                basis = "—"
                conclusion = "⚠️ 資料不足，需人工評估"

            results.append({
                "INCI 名稱": inci,
                "含量%": f"{conc:.5f}",
                "評估依據": basis,
                "NOAEL (mg/kg/day)": f"{noael_val:.2f}" if noael_val else "—",
                "DAp%": f"{dap:.0f}%",
                "Exposure (mg/kg bw/day)": f"{exposure:.6f}" if conc else "—",
                "SED (mg/kg/day)": "—" if cramer == "不需評估" else (f"{sed:.6f}" if conc else "—"),
                "PoD (mg/kg bw/day)": f"{noael_val:.2f}" if pod_source == "NOAEL" else "—",
                "ToE value": f"{pod_used:.7f}" if pod_source and pod_source.startswith("TTC") else "—",
                "Bioavailability%": f"{bioavail:.0f}%",
                "MoS": f"{mos:.1f}" if mos else "—",
                "RCR(MoE)": f"{rcr:.4f}" if rcr else "—",
                "PoD 來源": pod_source or "—",
                "結論": conclusion,
            })

        st.markdown("#### 計算結果")
        result_df = pd.DataFrame(results)
        st.dataframe(result_df, use_container_width=True, hide_index=True)

        # 儲存 NOAEL / DAp / Bioavailability / Cramer 覆寫值到資料庫
        updated_ings = db.get_pif_ingredients(pif_id)
        for i, row in edited_tox.iterrows():
            if i < len(updated_ings):
                noael_val = row["NOAEL (mg/kg/day)"]
                dap_val = row["DAp%"]
                bio_val = row["Bioavailability%"]
                cram_val = row["Cramer Class"]
                updated_ings[i]["noael_manual"] = float(noael_val) if pd.notna(noael_val) else None
                updated_ings[i]["dap_override"] = float(dap_val)
                updated_ings[i]["bioavailability_override"] = float(bio_val) if pd.notna(bio_val) else None
                updated_ings[i]["cramer_class"] = cram_val or None
        db.save_pif_ingredients(pif_id, updated_ings)

        # 儲存計算結果文字到章節資料
        summary_lines = []
        for r in results:
            summary_lines.append(
                f"- {r['INCI 名稱']}：含量 {r['含量%']}%，評估依據={r['評估依據']}，"
                f"NOAEL={r['NOAEL (mg/kg/day)']}，SED={r['SED (mg/kg/day)']}，MoS={r['MoS']}，{r['結論']}"
            )
        db.save_chapter_data(pif_id, 10, {
            "product_type": product_type_sel,
            "effective_daily_g": effective_daily_g,
            "results_summary": "\n".join(summary_lines),
            "results": results,
        })
        st.success("計算結果已儲存至第十章！")

    # 顯示上次計算結果
    saved = db.get_chapter_data(pif_id, 10)
    if saved.get("results"):
        st.divider()
        st.caption("上次計算結果（已儲存）：")
        st.dataframe(pd.DataFrame(saved["results"]), use_container_width=True, hide_index=True)


# ─── 第十六章：安全性評估匯總表（呈現層，讀第十章結果） ──────────────────────────

def render_chapter_16(db: IngredientDB, pif_id: int):
    st.subheader("十六、產品安全資料（安全評估報告）")
    st.markdown("#### 安全性評估匯總表")
    st.latex(r"MoS = \frac{PoD_{sys}}{SED} = \frac{PoD \times Bioavailability}{Exposure \times Dermal\ Absorption}")
    st.caption("PoDsys = PoD × 生體可用率；SED = 外用暴露量 × 皮膚吸收率；判定門檻 MoS ≥ 100；TTC 估算僅供輔助。")

    ch10 = db.get_chapter_data(pif_id, 10)
    if not ch10.get("results"):
        st.warning("⚠️ 尚未完成第十章 SED/MoS 計算，請先至第十章計算並儲存。")
    else:
        rows = [{
            "成分名稱": r.get("INCI 名稱", ""),
            "含量 w/w%": r.get("含量%", ""),
            "Exposure (mg/kg bw/day)": r.get("Exposure (mg/kg bw/day)", "—"),
            "皮膚吸收%": r.get("DAp%", ""),
            "SED (mg/kg bw/day)": r.get("SED (mg/kg/day)", "—"),
            "PoD (mg/kg bw/day)": r.get("PoD (mg/kg bw/day)", "—"),
            "ToE value": r.get("ToE value", "—"),
            "Bioavailability%": r.get("Bioavailability%", "—"),
            "MoS": r.get("MoS", "—"),
            "RCR(MoE)": r.get("RCR(MoE)", "—"),
            "MoE Approach": r.get("評估依據", "—"),
        } for r in ch10["results"]]
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

    st.divider()
    st.markdown("#### 安全評估結論與簽署")
    data = db.get_chapter_data(pif_id, 16)
    default = data.get("content", CHAPTER_TEMPLATES[16])
    content = st.text_area(
        "章節內容",
        value=default,
        height=350,
        key="ch16_content",
        label_visibility="collapsed",
    )
    if st.button("💾 儲存第十六章", type="primary", key="save_ch16"):
        db.save_chapter_data(pif_id, 16, {"content": content})
        st.success("第 16 章已儲存！")
        st.rerun()


# ─── 通用章節（文字填寫） ───────────────────────────────────────────────────────

CHAPTER_TEMPLATES = {
    2: "登錄確認文件編號：\n登錄日期：\n備註：",
    4: "□ 外包裝正面照片/圖示\n□ 外包裝背面照片/圖示\n□ 內容器照片/圖示\n□ 標籤（如有）\n□ 仿單/說明書（如有）\n\n標示符合性確認：\n□ 中文品名\n□ 製造廠名稱與地址\n□ 全成分名稱\n□ 內容量\n□ 保存期限或製造日期\n□ 批號或出廠日期\n□ 用途\n□ 用法用量\n□ 警語（如有）",
    5: "符合方式：\n□ 主管機關核發之 GMP 證書（證書號碼：　　，有效期限：　　）\n□ 第三方驗證機構核發之 GMP 證書\n□ 業者自我聲明書\n\n自我聲明書內容（如適用）：\n本公司聲明，[製造廠名稱] 位於 [地址] 之化粧品製造場所，遵循「化粧品優良製造準則」規定進行生產。\n聲明人：　　　　職稱：　　　　日期：",
    6: "製造流程概述：\n\n（請填入原料投入順序、溫度控制、混合時間、充填條件等）\n\n附件：\n□ 生產製造標準書（SOP）\n□ 批次生產及管制紀錄格式",
    7: "使用部位：\n建議用量：\n使用頻率：\n使用駐留時間：□ 沖洗型  □ 留置型\n適用族群：□ 一般成人  □ 嬰幼兒（3歲以下）  □ 老年人\n不適用族群/警語：\n可預見之誤用情況：\n\n使用說明（包裝標示文字）：",
    8: "不良反應管理制度建立日期：\n負責人員：\n\n不良反應記錄（上市前填「尚無」）：\n尚無不良反應紀錄。\n\n注意：發現嚴重不良反應時，應自得知之日起 15 日內向衛福部食藥署通報。",
    9: "成品規格：\n外觀：\n顏色：\n氣味：\n酸鹼值（pH）：\n黏度：\n比重：\n\n主要活性成分物化特性（分子量、溶解度、分配係數等）：",
    11: "試驗方法依據：□ ICH Q1A  □ ASEAN 準則  □ 其他：\n加速試驗條件：40°C / 75% RH，6個月\n長期試驗條件：25°C / 60% RH，12個月以上\n觀察頻率：T=0, 1M, 2M, 3M, 6M, 12M\n觀察項目：外觀、顏色、氣味、pH、黏度、微生物\n\n試驗結果：",
    12: "成品微生物規格：\n總生菌數：≤ 1000 CFU/g（一般產品）\n大腸桿菌：不得檢出\n金黃色葡萄球菌：不得檢出\n綠膿桿菌：不得檢出\n黴菌及酵母菌：≤ 100 CFU/g\n\n檢驗報告（請附上）：",
    13: "試驗方法：□ ISO 11930:2019  □ USP 37 Chapter <51>  □ 其他：\n\n防腐劑成分說明：\n\n試驗結果（請附上報告）：\n\n（如為微生物低風險產品，請說明豁免理由）",
    14: "功能宣稱：\n\n佐證資料：\n□ 推論功能之參考文獻\n□ 體外功能性評估報告（In vitro）\n□ 消費者實測報告（Consumer study）\n□ 特定成分含量檢驗報告\n□ 其他：\n\n注意：宣稱內容不得涉及醫療效能、虛偽或誇大。",
    15: "包裝材質清單：\n| 包裝部位 | 材質 | 容量/規格 | 供應商 |\n|---------|------|---------|--------|\n| 瓶身 | | | |\n| 瓶蓋 | | | |\n| 泵頭 | | | |\n\n包材相容性評估：\n□ 已進行包材安定性測試\n□ 已確認無遷移/溶出問題",
    16: "安全資料簽署人員資訊：\n姓名：\n學歷：\n職稱：\n服務單位：\n受訓證明：\n簽署日期：\n\n安全評估結論：\n本安全評估依據上述各項資料，就本產品中各成分之毒理特性、暴露量及成分相互作用進行評估，結論如下：\n1. 成分法規符合性：所有添加成分及濃度均符合「化粧品衛生安全管理法」及相關規定。\n2. 局部暴露評估：依使用部位、用量、頻率及留置時間計算，各成分之全身暴露量均在安全範圍內。\n3. 安全臨界值（MoS）：主要活性成分之 MoS 計算結果詳見第十章，MoS ≥ 100 為可接受。\n4. 微生物及化學品質：成品符合相關規格。\n5. 安定性：產品於建議儲存條件下安定性良好。\n6. 整體評估：本產品於正常及合理可預見之使用條件下，對消費者健康不會產生危害。\n\n聲明：\n經分析所有可取得之安全性資料，並根據當前科學知識據以結論，推定在預期正常合理使用條件下，本產品為可安全使用之產品，不致對人體健康造成傷害。\n\n簽名：____________  日期：____年__月__日",
}


NUM_ZH = {
    1: "一", 2: "二", 3: "三", 4: "四", 5: "五",
    6: "六", 7: "七", 8: "八", 9: "九", 10: "十",
    11: "十一", 12: "十二", 13: "十三", 14: "十四",
    15: "十五", 16: "十六",
}


def render_generic_chapter(db: IngredientDB, pif_doc: dict, chapter_num: int):
    pif_id = pif_doc["id"]
    name = CHAPTER_NAMES[chapter_num]
    st.subheader(f"{NUM_ZH[chapter_num]}、{name}")

    with st.expander("📋 本章節應包含的內容", expanded=False):
        st.info(CHAPTER_HINTS.get(chapter_num, ""))

    data = db.get_chapter_data(pif_id, chapter_num)
    default = data.get("content", CHAPTER_TEMPLATES.get(chapter_num, ""))

    content = st.text_area(
        "章節內容",
        value=default,
        height=350,
        key=f"ch{chapter_num}_content",
        label_visibility="collapsed",
    )

    if chapter_num in ATTACHMENT_CHAPTERS:
        st.divider()
        st.markdown("#### 📎 章節附件")
        render_chapter_attachments(db, pif_doc, chapter_num)
        st.divider()

    col1, col2 = st.columns([1, 5])
    with col1:
        if st.button(f"💾 儲存第{NUM_ZH[chapter_num]}章", type="primary", key=f"save_ch{chapter_num}"):
            saved = db.get_chapter_data(pif_id, chapter_num)
            saved["content"] = content
            db.save_chapter_data(pif_id, chapter_num, saved)
            st.success(f"第 {chapter_num} 章已儲存！")
            st.rerun()


# ─── 匯出區 ────────────────────────────────────────────────────────────────────

def render_export_panel(db: IngredientDB, pif_id: int, pif_doc: dict):
    st.divider()
    st.markdown("### 📤 匯出 PIF 文件")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📄 匯出 Word (.docx)", use_container_width=True):
            try:
                from pif_exporter import export_word
                out_path = Path("data") / f"{pif_doc['product_name']}_PIF.docx"
                export_word(pif_id, db, str(out_path))
                with open(out_path, "rb") as f:
                    st.download_button(
                        "⬇️ 下載 Word 檔",
                        data=f,
                        file_name=out_path.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
            except Exception as e:
                st.error(f"匯出失敗：{e}")
    with col2:
        if st.button("📋 匯出 PDF", use_container_width=True):
            try:
                from pif_exporter import export_word, export_pdf
                docx_path = Path("data") / f"{pif_doc['product_name']}_PIF.docx"
                pdf_path = Path("data") / f"{pif_doc['product_name']}_PIF.pdf"
                export_word(pif_id, db, str(docx_path))
                export_pdf(str(docx_path), str(pdf_path))
                with open(pdf_path, "rb") as f:
                    st.download_button(
                        "⬇️ 下載 PDF 檔",
                        data=f,
                        file_name=pdf_path.name,
                        mime="application/pdf",
                    )
            except Exception as e:
                st.error(f"匯出失敗：{e}")

    st.divider()
    st.markdown("#### 📎 原料 COA / SDS / IFRA 附錄")

    # 預覽缺檔警告
    file_rows = db.get_pif_material_coa_sds(pif_id)
    missing: list[str] = []
    seen_mids: set[int] = set()
    has_any_file = False
    for row in file_rows:
        mid = row["material_id"]
        if row["file_data"] is not None:
            has_any_file = True
        if mid not in seen_mids:
            seen_mids.add(mid)
            has_file = any(r["file_data"] is not None for r in file_rows if r["material_id"] == mid)
            if not has_file:
                code = row["ingredient_code"] or ""
                name = row["material_name"] or ""
                missing.append(f"{code} {name}".strip())

    if missing:
        st.warning(
            "以下原料尚無 COA/SDS/IFRA 原始檔，將在附錄索引中標記 ⚠️：\n"
            + "、".join(missing)
        )

    if st.button("📄 匯出 COA/SDS/IFRA 附錄 PDF", use_container_width=True, disabled=not has_any_file):
        try:
            from pif_exporter import export_coa_sds_appendix
            with st.spinner("合併 COA/SDS/IFRA 文件中..."):
                pdf_bytes, fname = export_coa_sds_appendix(pif_id, db)
            st.download_button(
                "⬇️ 下載 COA/SDS/IFRA 附錄",
                data=pdf_bytes,
                file_name=fname,
                mime="application/pdf",
                key="dl_coa_sds_appendix",
            )
        except Exception as e:
            st.error(f"匯出失敗：{e}")

    if not has_any_file:
        st.caption("尚無原料 COA/SDS/IFRA 原始檔。請至原料頁面重新上傳 PDF 文件後再使用此功能。")


# ─── 文件列表儀表板 ────────────────────────────────────────────────────────────

def render_document_list(db: IngredientDB):
    """顯示所有 PIF 文件的列表，提供開啟、刪除、新建操作。"""
    st.subheader("現有 PIF 文件")

    search_q = st.text_input("🔍 搜尋 PIF 文件", placeholder="輸入貨品編號或貨品名稱…", key="pif_search")

    with st.expander("➕ 新建 PIF 文件"):
        new_number = st.text_input("貨品編號（選填）", key="new_pif_number", placeholder="例：PIF-2026-001")
        new_name = st.text_input("貨品名稱 *", key="new_pif_name", placeholder="例：嬰幼潤澤爽身乳液")
        if st.button("建立", key="create_pif") and new_name.strip():
            new_id = db.create_pif_document(new_name.strip(), new_number.strip())
            st.session_state["selected_pif_id"] = new_id
            st.rerun()

    if st.button("📥 匯入 Excel", key="show_pif_xl_import"):
        st.session_state["show_pif_import"] = not st.session_state.get("show_pif_import", False)

    if st.session_state.get("show_pif_import"):
        with st.container(border=True):
            st.markdown("**匯入 Excel — 批次建立 PIF 文件**")
            st.caption(
                "Excel 欄位名稱包含以下關鍵字即可自動對應：\n"
                "- 貨品編號（或 code / no）\n"
                "- 貨品名稱（或 name）"
            )
            xl_file = st.file_uploader("選擇 Excel 檔案", type=["xlsx", "xls"], key="pif_xl_import")
            if xl_file:
                try:
                    df_xl = pd.read_excel(xl_file)
                    col_map: dict[str, str] = {}
                    for col in df_xl.columns:
                        cu = str(col).upper()
                        if ("編號" in col or "CODE" in cu or col.upper().endswith("NO")) and "code" not in col_map:
                            col_map["code"] = col
                        if ("名稱" in col or "NAME" in cu) and "name" not in col_map:
                            col_map["name"] = col

                    if "name" not in col_map:
                        st.error(f"找不到貨品名稱欄位，現有欄位：{list(df_xl.columns)}")
                    else:
                        st.caption(
                            f"欄位對應：貨品編號=`{col_map.get('code', '未偵測')}`、"
                            f"貨品名稱=`{col_map['name']}`"
                        )
                        st.dataframe(df_xl[[c for c in col_map.values() if c]].head(5),
                                     use_container_width=True)
                        st.caption(f"共 {len(df_xl)} 筆，預覽前 5 筆")

                        if st.button(f"✅ 確認匯入 {len(df_xl)} 筆", type="primary", key="confirm_pif_import"):
                            imported = 0
                            for _, row in df_xl.iterrows():
                                name = str(row.get(col_map["name"], "")).strip()
                                if not name or name.lower() == "nan":
                                    continue
                                code = str(row.get(col_map.get("code", ""), "")).strip()
                                db.create_pif_document(name, "" if code == "nan" else code)
                                imported += 1
                            st.success(f"已匯入 {imported} 筆 PIF 文件")
                            st.session_state["show_pif_import"] = False
                            st.rerun()
                except Exception as e:
                    st.error(f"讀取 Excel 失敗：{e}")

    docs = db.get_pif_documents()
    if not docs:
        st.info("尚無 PIF 文件，請點上方按鈕新建第一份文件。")
        return

    if search_q:
        docs = [
            d for d in docs
            if search_q.lower() in (d.get("document_number") or "").lower()
            or search_q.lower() in d["product_name"].lower()
        ]
        if not docs:
            st.info("找不到符合條件的 PIF 文件。")
            return

    completion_counts = db.get_all_pif_completion_counts()

    # 欄位標頭：文件編號 | 貨品名稱 | 章節進度 | 最後更新 | 操作
    hcols = st.columns([1.5, 3, 3, 1.5, 1.5])
    hcols[0].markdown("**貨品編號**")
    hcols[1].markdown("**貨品名稱**")
    hcols[2].markdown("**章節進度**")
    hcols[3].markdown("**最後更新**")
    hcols[4].markdown("**操作**")
    st.divider()

    for doc in docs:
        doc_id = doc["id"]
        done = min(completion_counts.get(doc_id, 0), 16)
        updated = (doc.get("updated_at") or "")[:10]
        confirm_key = f"confirm_del_{doc_id}"

        row = st.columns([1.5, 3, 3, 1.5, 1.5])
        row[0].write(doc.get("document_number") or "—")
        row[1].markdown(f"**{doc['product_name']}**")
        with row[2]:
            st.progress(done / 16)
            st.caption(f"{done}/16 章節完成")
        row[3].write(updated)
        with row[4]:
            btn_col1, btn_col2 = st.columns(2)
            if btn_col1.button("開啟", key=f"open_{doc_id}"):
                st.session_state["selected_pif_id"] = doc_id
                st.session_state.pop(confirm_key, None)
                st.rerun()
            if btn_col2.button("🗑️", key=f"del_{doc_id}", help="刪除此 PIF"):
                st.session_state[confirm_key] = True
                st.rerun()

        if st.session_state.get(confirm_key):
            st.warning(f"確定要刪除「{doc['product_name']}」嗎？此操作無法復原。")
            cfm_col1, cfm_col2 = st.columns([1, 5])
            if cfm_col1.button("確認刪除", key=f"cfm_{doc_id}"):
                _delete_pif_with_files(db, doc)
                st.session_state.pop(confirm_key, None)
                st.rerun()
            if cfm_col2.button("取消", key=f"cancel_{doc_id}"):
                st.session_state.pop(confirm_key, None)
                st.rerun()

        st.divider()


# ─── 主頁面 ────────────────────────────────────────────────────────────────────

def main():
    st.title("📄 PIF 文件製作")
    db = get_db()

    # 初始化模式狀態
    if "selected_pif_id" not in st.session_state:
        st.session_state["selected_pif_id"] = None

    # ── 列表模式 ─────────────────────────────────────────────────────────────────
    if st.session_state["selected_pif_id"] is None:
        render_document_list(db)
        return

    # ── 編輯模式 ─────────────────────────────────────────────────────────────────
    pif_id = st.session_state["selected_pif_id"]

    pif_doc = db.get_pif_document(pif_id)
    if not pif_doc:
        st.error("找不到此 PIF 文件。")
        st.session_state["selected_pif_id"] = None
        st.rerun()
        return

    # 返回列表按鈕
    if st.button("← 返回文件列表", key="back_to_list"):
        st.session_state["selected_pif_id"] = None
        st.session_state.pop(f"confirm_del_pif_{pif_id}", None)
        st.rerun()

    with st.expander("✏️ 編輯文件基本資訊", expanded=False):
        edit_col1, edit_col2 = st.columns(2)
        with edit_col1:
            new_product_name = st.text_input("貨品名稱", value=pif_doc.get("product_name", ""), key="edit_product_name")
        with edit_col2:
            new_doc_num = st.text_input("貨品編號", value=pif_doc.get("document_number", ""), key="edit_doc_number")
        if st.button("💾 儲存基本資訊", key="save_basic_info"):
            updates: dict = {}
            if new_product_name.strip():
                updates["product_name"] = new_product_name.strip()
            updates["document_number"] = new_doc_num.strip()

            old_dir = pif_upload_dir(db, pif_doc)
            db.update_pif_document(pif_id, updates)
            new_dir = pif_upload_dir(db, db.get_pif_document(pif_id))

            if new_dir != old_dir and old_dir.exists():
                if new_dir.exists():  # 目標已存在：逐檔搬入
                    for src in old_dir.iterdir():
                        shutil.move(str(src), str(new_dir / src.name))
                    old_dir.rmdir()
                else:
                    new_dir.parent.mkdir(parents=True, exist_ok=True)
                    old_dir.rename(new_dir)
                _rewrite_chapter_paths(db, pif_id, old_dir, new_dir)
            st.rerun()

    doc_num_display = pif_doc.get("document_number") or "—"
    st.subheader(f"📄 {pif_doc['product_name']}")
    st.caption(f"編號：{doc_num_display} ｜ 最後更新：{pif_doc.get('updated_at', '')[:10]}")

    with st.sidebar:
        st.markdown("## 章節進度")
        statuses = db.get_chapter_completion_status(pif_id)
        done = sum(1 for v in statuses.values() if v)
        st.progress(done / 16, text=f"{done}/16 章節完成")
        st.divider()

        chapters = list(CHAPTER_NAMES.keys())
        # 目前章節存在 DB（pif_documents.current_chapter）。DB 讀取跟 websocket / session
        # 狀態無關，每個 run 都必定讀得到，撐得過 AI 解析等長操作造成的重連 / session 重置。
        _seed = pif_doc.get("current_chapter") or 1
        if _seed not in chapters:
            _seed = 1

        def _save_current_chapter():
            db.set_current_chapter(pif_id, st.session_state["chapter_selector"])

        # 每個 run 開頭先把 widget 值對齊 DB。使用者點選時 on_change 已先把新值寫進 DB，
        # 故這裡不會蓋掉操作；重連 / session 清空後也能靠 DB 的值強制回到正確章節。
        st.session_state["chapter_selector"] = _seed

        chosen_chapter = st.radio(
            "選擇章節",
            chapters,
            format_func=lambda n: f"{chapter_emoji(statuses.get(n, False))} 第{n}章 {CHAPTER_NAMES[n][:10]}",
            key="chapter_selector",
            on_change=_save_current_chapter,
        )

        st.divider()
        confirm_del_key = f"confirm_del_pif_{pif_id}"
        if st.session_state.get(confirm_del_key):
            st.warning(f"確定要刪除「{pif_doc['product_name']}」嗎？此操作無法復原。")
            if st.button("✔️ 確認刪除", key="cfm_del_pif", type="primary"):
                _delete_pif_with_files(db, pif_doc)
                st.session_state.pop(confirm_del_key, None)
                st.session_state["selected_pif_id"] = None
                st.rerun()
            if st.button("取消", key="cancel_del_pif"):
                st.session_state.pop(confirm_del_key, None)
                st.rerun()
        elif st.button("🗑️ 刪除此 PIF", key="del_pif"):
            st.session_state[confirm_del_key] = True
            st.rerun()

    chapter_num = chosen_chapter

    if chapter_num == 1:
        render_chapter_1(db, pif_id)
    elif chapter_num == 2:
        render_chapter_2(db, pif_doc)
    elif chapter_num == 3:
        render_chapter_3(db, pif_doc)
    elif chapter_num == 10:
        render_chapter_10(db, pif_id)
    elif chapter_num == 16:
        render_chapter_16(db, pif_id)
    else:
        render_generic_chapter(db, pif_doc, chapter_num)

    render_export_panel(db, pif_id, pif_doc)


if __name__ == "__main__":
    main()
else:
    main()
